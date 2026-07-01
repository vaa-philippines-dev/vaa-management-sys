from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

t = doc.add_heading('RSC Serialization Bug Hunt \u2014 Fix Log', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('June 30, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

doc.add_heading('Overview', 1)
doc.add_paragraph(
    'After deploying the admin panel to production (Vercel), the /admin page failed with '
    'error code 1621801304: "Functions cannot be passed directly to Client Components unless '
    'you explicitly expose it by marking it with \'use server\'. Or maybe you meant to call '
    'this function rather than return it."'
)
doc.add_paragraph(
    'The error object showed {$$typeof: ..., render: function, displayName: ...} \u2014 '
    'indicating a function was present in the React Server Component (RSC) serialization payload. '
    'This triggered a chain of 9 fixes over several hours to root out every instance of '
    'function-passing across the Server/Client boundary.'
)
doc.add_paragraph(
    'This document catalogues each fix in chronological order, the affected files, root cause, '
    'and the technical rationale behind the solution.'
)

doc.add_heading('Technical Background', 1)
doc.add_paragraph(
    'In Next.js App Router, Server Components render on the server and their output is serialized '
    'into a JSON-like payload sent to the client. Client Components (marked with "use client") '
    'hydrate from this payload. The serialization boundary does not support functions \u2014 any '
    'function reference crossing from Server to Client causes the digest 1621801304 error.'
)
doc.add_paragraph('Three distinct categories of function-leak were discovered:')

points = [
    '@base-ui/react library internals \u2014 components using useRender hook inject a render property '
    '(a function) into their RSC output, which Next.js tries to serialize.',
    'Inline arrow functions used as form action handlers \u2014 Next.js captures the closure and attempts '
    'to serialize the parent scope (including all its functions).',
    'Component references (lucide-react icons, React.ComponentType props) passed from Server Components '
    'to Client Components \u2014 these are function references and cannot be serialized.',
]
for p_text in points:
    doc.add_paragraph(p_text, style='List Bullet')

doc.add_heading('Fix #1 \u2014 DepartmentCard bound server action', 1)
doc.add_heading('Commit: cf51950', 2)
doc.add_paragraph(
    'File: components/admin/DepartmentCard.tsx'
)
doc.add_paragraph(
    'Root Cause: The edit form used an anonymous inline arrow function as its action handler. '
    'When Next.js serializes a Client Component that uses inline functions as server action '
    'closures, it captures the entire parent scope \u2014 including all function references in the module.'
)
doc.add_paragraph(
    'Fix: Replaced the inline arrow function with editDepartment.bind(null, dept.id), binding '
    'the imported server action to the department ID. This avoids creating a closure that captures '
    'the module scope.'
)
doc.add_heading('Code Diff', 2)
doc.add_paragraph(
    'Before: action={async (fd) => { "use server"; await editDepartment(dept.id, fd) }}\n'
    'After:  const editAction = editDepartment.bind(null, dept.id); action={editAction}'
)

doc.add_heading('Fix #2 \u2014 DialogTrigger render prop', 1)
doc.add_heading('Commit: 0754b2d', 2)
doc.add_paragraph(
    'File: VAProfileEditor.tsx'
)
doc.add_paragraph(
    'Root Cause: The @base-ui/react Dialog.Trigger component accepted a render prop with complex '
    'JSX children. The render prop is a function; Base UI passes it through as a render: function '
    'property in the RSC payload, and Next.js fails to serialize it.'
)
doc.add_paragraph(
    'Fix: Replaced the render prop pattern with a controlled open/close state using useState. '
    'The trigger button now uses onClick to toggle visibility rather than relying on Base UI\u2019s '
    'internal render function propagation.'
)

doc.add_heading('Fix #3 \u2014 DialogClose render prop', 1)
doc.add_heading('Commit: a7560d0', 2)
doc.add_paragraph(
    'File: components/vas/VAProfileEditor.tsx'
)
doc.add_paragraph(
    'Root Cause: Dialog.Close was used with render={<Button>...</Button>}. The render prop '
    'contains a displayName (function identifier) that fails RSC serialization for the same reason \u2014 '
    'Base UI converts render props into function references in the RSC tree.'
)
doc.add_paragraph(
    'Fix: Replaced render={<Button>} with a className-based approach, avoiding the function '
    'reference entirely.'
)

doc.add_heading('Fix #4 \u2014 Inline form actions in admin/users', 1)
doc.add_heading('Commit: f7b83df', 2)
doc.add_paragraph(
    'Files: app/(dashboard)/admin/users/page.tsx, app/(dashboard)/admin/users/actions.ts'
)
doc.add_paragraph(
    'Root Cause: Several form handlers used inline arrow functions with "use server" directives. '
    'Each inline function captured the surrounding server component\u2019s scope, including database '
    'query results and Prisma instances, causing RSC serialization failures when the bound forms '
    'were serialized for client hydration.'
)
doc.add_paragraph(
    'Fix: Extracted all form action logic into exported server functions in actions.ts '
    '(updateUserRoleByForm, updateUserTypeByForm, assignDeptByForm, assignTempRoleByForm). '
    'Forms now reference these exported functions via .bind() instead of inline closures.'
)

doc.add_heading('Fix #5 \u2014 @base-ui Dialog replaced with custom Modal', 1)
doc.add_heading('Commit: d58b547', 2)
doc.add_paragraph(
    'File: components/vas/VAProfileEditor.tsx'
)
doc.add_paragraph(
    'Root Cause: Despite fixing render props, the @base-ui Dialog primitive itself contained '
    'internal render function properties in its SSR output. The library\u2019s architecture fundamentally '
    'relies on render functions that break the RSC serialization contract.'
)
doc.add_paragraph(
    'Fix: Completely removed @base-ui/react Dialog. Replaced with a custom Modal component '
    'using native HTML dialog element and CSS transitions, with no render function references.'
)

doc.add_heading('Fix #6 \u2014 Badge/Button @base-ui primitives in admin page', 1)
doc.add_heading('Commit: 56485a5', 2)
doc.add_paragraph(
    'Files: app/(dashboard)/admin/page.tsx, components/admin/DepartmentCard.tsx'
)
doc.add_paragraph(
    'Root Cause: After fixing Dialog, the admin page still failed because @base-ui/react Badge '
    'and Button components use useRender internally. This hook injects a render property into '
    'every RSC tree node that uses these components, creating function references wherever Badge '
    'or Button are rendered in Server Components.'
)
doc.add_paragraph(
    'Fix: Replaced all Badge and Button uses in admin/page.tsx and DepartmentCard.tsx with '
    'plain HTML elements styled with the same Tailwind classes. Eliminated the Badge/Button '
    'component imports from @base-ui in the admin context.'
)

doc.add_heading('Fix #7 \u2014 Global Badge/Button replacement with plain HTML', 1)
doc.add_heading('Commit: ca19fe8', 2)
doc.add_paragraph(
    'Files: components/ui/badge.tsx, components/ui/button.tsx'
)
doc.add_paragraph(
    'Root Cause: Fix #6 was only applied to the admin page files. However, Badge and Button had '
    '@base-ui internal render properties that could fail anywhere in the application where these '
    'components rendered within a Server Component tree.'
)
doc.add_paragraph(
    'Fix: Rewrote components/ui/badge.tsx and components/ui/button.tsx to use plain HTML '
    'elements (\u003Cdiv\u003E, \u003Cspan\u003E) instead of @base-ui primitives. Maintained all class variance '
    'authority (CVA) styling and accessibility attributes. These are now pure presentational '
    'components with zero function references in their RSC output.'
)

doc.add_heading('Fix #8 \u2014 Input/ScrollArea + remaining inline actions', 1)
doc.add_heading('Commit: 3ddac30', 2)
doc.add_paragraph(
    'Files: components/ui/input.tsx, components/ui/scroll-area.tsx, '
    'app/(dashboard)/admin/users/actions.ts, app/(dashboard)/admin/users/page.tsx'
)
doc.add_paragraph(
    'Root Cause: Input and ScrollArea still imported from @base-ui/react (Input had Base UI '
    'wrappers; ScrollArea used Base UI ScrollArea primitives). Additionally, the admin users page '
    'still had some inline server action patterns.'
)
doc.add_paragraph(
    'Fix: Replaced Base UI Input with native \u003Cinput\u003E styled with Tailwind. Replaced Base UI '
    'ScrollArea with a native overflow-based implementation. Added the remaining bound action '
    'exports (toggleUserActive bind pattern) to actions.ts.'
)

doc.add_heading('Fix #9 \u2014 Delete @base-ui/react entirely', 1)
doc.add_heading('Commit: 7fad3a0', 2)
doc.add_paragraph(
    'Files: components/ui/alert-dialog.tsx, components/ui/dialog.tsx, components/ui/select.tsx, '
    'components/ui/tabs.tsx, package.json, package-lock.json'
)
doc.add_paragraph(
    'Root Cause: Even after rewriting all @base-ui usages in the app code, the import statements '
    'in unused UI components (alert-dialog, dialog, select, tabs) could still be bundled and cause '
    'issues. The package\u2019s internal architecture is fundamentally incompatible with RSC.'
)
doc.add_paragraph(
    'Fix: Deleted all 4 remaining @base-ui component files. Removed @base-ui/react from '
    'package.json dependencies and package-lock.json. Ran npm install to clean up. Total removal: '
    '763 lines deleted across 6 files.'
)

doc.add_heading('Fix #10 \u2014 Icon component passed to Client Component', 1)
doc.add_heading('Commit: b512716', 2)
doc.add_paragraph(
    'Files: app/(dashboard)/admin/page.tsx, components/admin/DepartmentCard.tsx'
)
doc.add_paragraph(
    'Root Cause: After all @base-ui fixes, the admin page still failed. The DepartmentCard '
    '(a Client Component with "use client") received an icon prop of type React.ComponentType \u2014 '
    'a reference to a lucide-react icon function (Shield, Users, Building2, etc.). Server Components '
    'cannot pass function references to Client Components.'
)
doc.add_paragraph(
    'Fix: Changed the admin page to pre-render the icon as JSX on the server side: '
    'icon={\u003CIcon className="h-4 w-4 text-muted-foreground" /\u003E} instead of icon={Icon}. '
    'Changed DepartmentCard\u2019s icon prop type from React.ComponentType to React.ReactNode and '
    'rendered it as {icon} instead of \u003CIcon className="..." /\u003E.'
)
doc.add_heading('Code Diff', 2)
doc.add_paragraph(
    'Server (admin/page.tsx):\n'
    '  Before: \u003CDepartmentCard ... icon={Icon} /\u003E\n'
    '  After:  \u003CDepartmentCard ... icon={\u003CIcon className="..." /\u003E} /\u003E\n\n'
    'Client (DepartmentCard.tsx):\n'
    '  Before: icon: React.ComponentType\u003C{ className?: string }\u003E  ...  \u003CIcon className="..." /\u003E\n'
    '  After:  icon: React.ReactNode  ...  {icon}'
)

doc.add_heading('Root Cause Summary', 1)

table = doc.add_table(rows=11, cols=4)
table.style = 'Table Grid'
for j, h in enumerate(['#', 'Category', 'Pattern', 'Resolution']):
    table.rows[0].cells[j].text = h
    for p in table.rows[0].cells[j].paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ['1-3', 'Inline Closures', 'Arrow functions in form action', 'Export to actions.ts + .bind()'],
    ['4-5', 'Base UI render prop', 'render={\u003CComponent\u003E}', 'Controlled state / className'],
    ['6-8', 'Base UI internals', 'useRender injects render: function', 'Plain HTML replacements'],
    ['9', 'Base UI package', 'Library fundamentally incompatible', 'Uninstall + delete all files'],
    ['10', 'Component as prop', 'React.ComponentType passed to Client', 'Pre-render as ReactNode on server'],
]

for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_heading('Key Takeaway', 1)
doc.add_paragraph(
    'The fundamental rule is: no function (including React components, callbacks, or render props) '
    'may cross the Server Component \u2192 Client Component boundary. This includes:'
)
takeaways = [
    'React component references (lucide icons, CustomComponent) passed as props',
    'Arrow functions used as form action handlers in Client Components',
    'Library internals that use render functions (e.g., @base-ui useRender)',
    'Any prop typed as React.ComponentType or () =\u003E void on a "use client" component',
]
for t_text in takeaways:
    doc.add_paragraph(t_text, style='List Bullet')

doc.add_paragraph(
    '\nThe correct pattern is to pre-render functions on the server side into serializable values '
    '(ReactNode, string, number, plain objects) before passing them to Client Components.'
)

doc.add_heading('Files Affected Across All Fixes', 1)
files_affected = [
    'app/(dashboard)/admin/page.tsx \u2014 Badge/Button replacement, icon ReactNode fix',
    'app/(dashboard)/admin/users/page.tsx \u2014 Bound server actions, inline removal',
    'app/(dashboard)/admin/users/actions.ts \u2014 New exported action wrappers',
    'components/admin/DepartmentCard.tsx \u2014 Bound actions, HTML replacements, icon type fix',
    'components/vas/VAProfileEditor.tsx \u2014 Dialog to Modal, render prop removal',
    'components/ui/badge.tsx \u2014 @base-ui Badge \u2192 plain HTML',
    'components/ui/button.tsx \u2014 @base-ui ButtonPrimitive \u2192 plain HTML',
    'components/ui/input.tsx \u2014 @base-ui Input \u2192 native \u003Cinput\u003E',
    'components/ui/scroll-area.tsx \u2014 @base-ui ScrollArea \u2192 native overflow',
    'components/ui/alert-dialog.tsx \u2014 Deleted (@base-ui)',
    'components/ui/dialog.tsx \u2014 Deleted (@base-ui)',
    'components/ui/select.tsx \u2014 Deleted (@base-ui)',
    'components/ui/tabs.tsx \u2014 Deleted (@base-ui)',
    'package.json \u2014 @base-ui/react removed',
    'package-lock.json \u2014 Cleaned',
]
for f in files_affected:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Verification', 1)
doc.add_paragraph(
    'After Fix #10, the Next.js production build completed successfully:\n'
    '\u2022 All 21 pages compiled without errors\n'
    '\u2022 TypeScript check passed\n'
    '\u2022 /admin route generated as a dynamic (\u0192) server-rendered page\n'
    '\u2022 Deployed to Vercel via git push (commit b512716)'
)

p = doc.add_paragraph('\n\u2014 End of Document \u2014')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('RSC-Serialization-Bug-Hunt.docx')
print('Saved RSC-Serialization-Bug-Hunt.docx')

