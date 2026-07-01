from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

t = doc.add_heading('201 File Upload \u2014 Google Drive Integration', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Feature Completion Documentation').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('June 30, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

doc.add_heading('Overview', 1)
doc.add_paragraph(
    'The 201 File Upload feature enables HR and admin users to upload official 201 documents '
    '(Passport Photo, PhilHealth Photo, Signed Contract) for each Virtual Assistant directly to '
    'a Google Shared Drive. Files are organized in a nested folder structure, the upload includes '
    'real-time progress tracking, toast notifications with a "View Document" link, and visual '
    'highlighting of newly uploaded documents on the VA profile card.'
)

doc.add_heading('Architecture', 1)

doc.add_heading('Component Layers', 2)
layers = [
    'Upload Route: app/api/upload/route.ts \u2014 Next.js API endpoint that handles multipart file upload, creates/uses nested Google Drive folders, stores the file, and upserts the user profile with the Drive URL.',
    'FileUploadField: components/vas/FileUploadField.tsx \u2014 Standalone client component for single-file upload with progress bar, error handling, and "View uploaded file" link.',
    'Files201Content + UploadRow: components/vas/VAProfileEditor.tsx \u2014 Modal dialog content with three upload rows (Passport, PhilHealth, Contract), each with XHR-based upload, progress percentage, toast notification, and "Just uploaded" pulse badge.',
    'DocBadge: components/vas/VAProfileEditor.tsx \u2014 Preview badge on the VA profile card showing upload status with icons. Highlights with a green ring animation for 8 seconds after a fresh upload.',
    'Google Drive Library: lib/google/drive.ts \u2014 Shared utility for listing Drive files, creating folders, and uploading files using googleapis SDK with service account auth.',
    'Debug Endpoint: app/api/debug-drive/route.ts \u2014 Diagnostic endpoint that lists shared drives, parent folder contents, and detects duplicate folder names.',
]
for layer in layers:
    doc.add_paragraph(layer, style='List Bullet')

doc.add_heading('Google Drive Folder Structure', 2)
doc.add_paragraph(
    'Files are uploaded into a nested hierarchy under the configured Shared Drive parent folder:'
)
doc.add_paragraph(
    'GOOGLE_DRIVE_PARENT_FOLDER_ID (Shared Drive root)\n'
    '  \u2514 201 VA | [First Last]\n'
    '       \u2514 Passport\n'
    '       \u2514 Philhealth\n'
    '       \u2514 Profile Picture  (for signed contracts)'
)
doc.add_paragraph(
    'The folder names are determined by the DOC_TYPE_FOLDERS mapping:\n'
    '  passportPhoto \u2192 "Passport"\n'
    '  philhealthPhoto \u2192 "Philhealth"\n'
    '  signedContract \u2192 "Profile Picture"'
)
doc.add_paragraph(
    'Folders are created on-demand via findOrCreateFolder() \u2014 it checks if a folder with '
    'the given name already exists under the parent before creating a new one, avoiding duplicates.'
)

doc.add_heading('Upload Flow', 1)

steps = [
    ('1. User selects a file',
     'In the 201 Files modal (Files201Content), the UploadRow component renders a hidden '
     '\u003Cinput type="file"\u003E triggered by a styled label. On file selection, handleFile() is called '
     'with the File object.'),
    ('2. XHR upload begins',
     'A FormData object is constructed with file, vaName, fieldName, and profileId. The upload '
     'uses XMLHttpRequest (not fetch) to enable real-time progress tracking via the '
     'xhr.upload "progress" event. Progress updates in percentage are shown inline with a '
     'blue progress bar and percentage text.'),
    ('3. Server processes upload',
     'The POST /api/upload endpoint:\n'
     '  a. Authenticates with Google service account (drive scope)\n'
     '  b. Verifies the configured parent folder is accessible (supportsAllDrives: true)\n'
     '  c. Finds or creates the VA folder: "201 VA | {fullName}"\n'
     '  d. Finds or creates the document type subfolder (Passport/Philhealth/Profile Picture)\n'
     '  e. Uploads the file with sanitized filename\n'
     '  f. Updates UserProfile in database with the Drive webViewLink'),
    ('4. Toast notification',
     'On success, a toast notification appears with the message "[Label] uploaded successfully!" '
     'and an action button "View Document" that opens the Drive URL in a new tab. On failure, '
     'error details are shown in a red toast.'),
    ('5. Green highlight animation',
     'The UploadRow border and background turn green with a ring animation, and a "Just uploaded" '
     'pulse badge appears for 8 seconds. The parent DocBadge on the profile card also highlights '
     'with a green border, ring, and glow background for 8 seconds.'),
    ('6. Save to database',
     'The "Save Changes" button at the bottom of the modal triggers updateUserProfileFiles() '
     'which persists all three URLs (passportPhoto, philhealthPhoto, signedContract) to the '
     'UserProfile record via Prisma.'),
]
for title, desc in steps:
    doc.add_heading(title, 2)
    doc.add_paragraph(desc)

doc.add_heading('Key Technical Components', 1)

doc.add_heading('FileUploadField (components/vas/FileUploadField.tsx)', 2)
doc.add_paragraph(
    'A standalone reusable Client Component that handles single-file upload to /api/upload.\n'
    'Accepts: label, currentUrl, fieldName, vaName, profileId, onUploaded callback.\n'
    'Uses XMLHttpRequest for real-time progress (percentage + progress bar).\n'
    'Displays: upload button, progress during upload, "View uploaded file" link on success,\n'
    'error state with clear button, and "No file uploaded" placeholder.'
)

doc.add_heading('UploadRow in VAProfileEditor (lines 550-672)', 2)
doc.add_paragraph(
    'Same XHR-based upload logic as FileUploadField, plus:\n'
    '  \u2022 Toast notification with "View Document" action button via sonner\n'
    '  \u2022 Green highlight: border + ring animation + "Just uploaded" pulse badge (8 sec)\n'
    '  \u2022 Error handling with AlertCircle icon and red text\n'
    '  \u2022 Styled as a bordered card row with icon, label, file input, and progress area'
)

doc.add_heading('DocBadge (lines 299-338)', 2)
doc.add_paragraph(
    'A small badge component shown on the VA profile card preview for each 201 document.\n'
    'Displays: icon (IdCard/Camera/FileText), label (Passport/PhilHealth/Contract), and\n'
    'a "View \u2197" link if a URL exists, or "Missing" text if not.\n'
    'When highlighted=true: green border (#22c55e) + ring outline + green glow background\n'
    'with transition duration 300ms. Resets after 8 seconds via handleRecentUpload().'
)

doc.add_heading('API Route (app/api/upload/route.ts)', 2)
doc.add_paragraph(
    'POST /api/upload\n'
    '  Form fields: file, vaName, fieldName, profileId\n'
    '  Returns: { success, url, field, folder, fullPath }\n\n'
    'Caches root folder ID in memory after first verification.\n'
    'Uses supportsAllDrives: true on all Drive API calls for Shared Drive compatibility.\n'
    'Sanitizes file names: replaces non-word chars with underscores.\n'
    'Upserts UserProfile record: creates profile if missing, updates field if exists.\n'
    'In development, returns stack trace in error response.'
)

doc.add_heading('Google Drive Auth (lib/google/drive.ts)', 2)
doc.add_paragraph(
    'Provides reusable helpers: listDriveFiles(), createDriveFolder(), uploadFileToDrive().\n'
    'Uses googleapis SDK with service account credentials from env vars:\n'
    '  GOOGLE_SERVICE_ACCOUNT_EMAIL\n'
    '  GOOGLE_PRIVATE_KEY (with \\n replacement for multiline keys)\n'
    '  GOOGLE_DRIVE_PARENT_FOLDER_ID\n'
    'Scope: https://www.googleapis.com/auth/drive (full drive, not drive.file)\n'
    'Required because service accounts need full drive scope to access Shared Drive folders\n'
    'they did not create. Root folder ID is cached in module-level variable after first access.'
)

doc.add_heading('Debug Endpoint (app/api/debug-drive/route.ts)', 2)
doc.add_paragraph(
    'GET /api/debug-drive \u2014 Diagnostic tool for troubleshooting Drive access.\n'
    'Returns: service account email, shared drives list, parent folder metadata,\n'
    'folder children (names + IDs), child count, and duplicate folder name detection.\n'
    'Used during development to verify service account permissions, folder structure, and\n'
    'identify potential duplicate 201 VA folders.'
)

doc.add_heading('Bug Fixes & Iterations', 1)

fixes = [
    ('Service account scope',
     'Initially used drive.file scope, but service accounts need drive scope to access '
     'Shared Drive folders they did not create. Fixed in 5cdabd4.'),
    ('Service account root fallback',
     'Attempted to fall back to service account root folder. Removed because service accounts '
     'have no storage quota for personal root. Fixed in d485739.'),
    ('Shared Drive API calls',
     'Added supportsAllDrives: true to all drive.files.get/list/create calls. Without this, '
     'Shared Drive folders return 404. Fixed in 43566b4.'),
    ('camelCase field names',
     'Prisma schema uses camelCase (passportPhoto, philhealthPhoto, signedContract). '
     'Database upsert needed matching field names. Fixed in 039d04c.'),
    ('Nested folder structure',
     'Moved from flat uploads into "201 VA | Name / DocType" nested structure. Added '
     'findOrCreateFolder() to avoid duplicates. Implemented in 90ed34b.'),
    ('Toast notification + highlight',
     'Added sonner toast with "View Document" action link and green highlight animation '
     'on DocBadge and UploadRow for 8 seconds after upload. Implemented in 90ed34b.'),
    ('Duplicate folder detection',
     'debug-drive endpoint now lists parent folder children and detects duplicate folder '
     'names by counting occurrences. Implemented in db13d51.'),
]
for title, desc in fixes:
    doc.add_heading(title, 2)
    doc.add_paragraph(desc)

doc.add_heading('Database Schema', 1)

doc.add_heading('UserProfile (used for 201 file URLs)', 2)
doc.add_paragraph(
    'The upload API upserts UserProfile records with these fields:\n'
    '  passportPhoto: String? \u2014 Google Drive URL for passport photo\n'
    '  philhealthPhoto: String? \u2014 Google Drive URL for PhilHealth photo\n'
    '  signedContract: String? \u2014 Google Drive URL for signed contract'
)

doc.add_heading('VADocument (future use)', 2)
doc.add_paragraph(
    'model VADocument {\n'
    '  id, vaProfileId, documentType (DocumentType enum), fileName,\n'
    '  googleDriveUrl, mimeType, fileSize, uploadedBy, uploader,\n'
    '  expiresAt, notes, createdAt\n'
    '}\n\n'
    'This model exists for future document management beyond 201 files. The current 201 file '
    'implementation stores URLs directly on UserProfile for simplicity. Additional document '
    'types (CONTRACT, GOVERNMENT_ID, NDA, CLEARANCE, etc.) can use VADocument going forward.'
)

doc.add_heading('Environment Variables Required', 1)
env_vars = [
    ('GOOGLE_SERVICE_ACCOUNT_EMAIL', 'Service account email from Google Cloud Console'),
    ('GOOGLE_PRIVATE_KEY', 'Service account private key (PEM format, with \u005Cn newlines)'),
    ('GOOGLE_DRIVE_PARENT_FOLDER_ID', 'Shared Drive folder ID where 201 folders are created'),
]
table = doc.add_table(rows=len(env_vars) + 1, cols=2)
table.style = 'Table Grid'
for j, h in enumerate(['Variable', 'Description']):
    table.rows[0].cells[j].text = h
    for p in table.rows[0].cells[j].paragraphs:
        for r in p.runs:
            r.bold = True
for i, (var, desc) in enumerate(env_vars):
    table.rows[i + 1].cells[0].text = var
    table.rows[i + 1].cells[1].text = desc

doc.add_heading('UI States Covered', 1)
states = [
    'Empty: "No file uploaded" placeholder text in muted gray',
    'Uploading: File name spinner + progress bar (percentage + animated fill)',
    'Success: Green checkmark + "View file" link to Drive URL',
    'Error: Red AlertCircle icon + error message + X clear button',
    'Replace: Uploaded files show "Replace" button (not "Upload")',
    'Highlight: 8-second green ring animation on card + pulse badge "Just uploaded"',
    'Loading state: Progress bar with percentage text during XHR transfer',
]
for s in states:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Files Changed Across All Drive/201 Commits', 1)
files = [
    'app/api/upload/route.ts \u2014 Main upload endpoint (129 lines)',
    'app/api/debug-drive/route.ts \u2014 Drive diagnostic endpoint (66 lines)',
    'lib/google/drive.ts \u2014 Drive utility library (124 lines)',
    'components/vas/VAProfileEditor.tsx \u2014 201 Files modal + UploadRow + DocBadge (676 lines)',
    'components/vas/FileUploadField.tsx \u2014 Standalone file upload component (147 lines)',
    'prisma/schema.prisma \u2014 VADocument model + DocumentType enum additions',
    'app/(dashboard)/vas/[id]/page.tsx \u2014 VA profile page with upload integration',
    'app/(dashboard)/vas/actions.ts \u2014 Server actions for VA profile updates',
]
for f in files:
    doc.add_paragraph(f, style='List Bullet')

p = doc.add_paragraph('\n\u2014 End of Document \u2014')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('201-File-Upload-Google-Drive.docx')
print('Saved 201-File-Upload-Google-Drive.docx')

