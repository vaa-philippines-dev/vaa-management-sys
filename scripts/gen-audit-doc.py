from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ──
t = doc.add_heading('AUDIT LOGS MODULE\nDocumentation', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('July 1, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── Module Overview ──
doc.add_heading('MODULE OVERVIEW', 1)
doc.add_paragraph(
    'The Audit Logs Module provides comprehensive tracking of every significant action performed '
    'within the VAA Management System. Every user creation, role change, department assignment, '
    'file upload, and status toggle is automatically logged with full context — who did it, what '
    'was changed, the before and after values, and when it happened.'
)
doc.add_paragraph(
    'The module serves three purposes:\n'
    '1. Accountability — Every admin action is traceable to the specific user who performed it\n'
    '2. Debugging — When something goes wrong, audit logs show exactly what changed and when\n'
    '3. Compliance — Provides an immutable record of all system modifications for audits'
)
doc.add_paragraph(
    'Audit logs are written automatically by the system — no user action is required to enable '
    'logging. The logs are viewable by Super Admins and System Admins through a dedicated '
    'viewer page with filtering and search capabilities.'
)

# ── Field Summary ──
doc.add_heading('FIELD SUMMARY', 1)

fields_data = [
    ['1', 'Actor', 'The user who performed the action', 'Lookup (User)', 'Yes', 'No'],
    ['2', 'Action', 'Type of action performed', 'Enum (14 values)', 'Yes', 'No'],
    ['3', 'Entity Type', 'What was affected (User, Department, etc.)', 'Text', 'Yes', 'No'],
    ['4', 'Entity ID', 'The specific record identifier', 'Text', 'Yes', 'No'],
    ['5', 'Old Values', 'Properties before the change', 'JSON Object', 'No', 'No'],
    ['6', 'New Values', 'Properties after the change', 'JSON Object', 'No', 'No'],
    ['7', 'Metadata', 'Additional context about the action', 'JSON Object', 'No', 'No'],
    ['8', 'Department', 'Department scope (if applicable)', 'Lookup (Department)', 'No', 'No'],
    ['9', 'Date & Time', 'When the action occurred', 'DateTime (auto)', 'Yes (auto)', 'No'],
]

table = doc.add_table(rows=len(fields_data) + 1, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['#', 'Field Name', 'Description', 'Data Type', 'Required', 'Unique']):
    cell = table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(fields_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

# ── Action Types ──
doc.add_heading('ACTION TYPES', 1)
doc.add_paragraph(
    'Every audit log entry is classified by one of 14 action types. Each type has a color '
    'in the viewer for quick visual identification:'
)

action_data = [
    ['CREATE', 'Something was created', 'User account, department, temp role'],
    ['UPDATE', 'Something was modified', 'User type, department name, profile fields'],
    ['DELETE', 'Something was removed', 'Membership, temporary role'],
    ['STATUS_CHANGE', 'Active/inactive toggled', 'User enabled/disabled, department activation'],
    ['TRANSFER', 'Something was transferred', 'Reserved for future use'],
    ['APPROVE', 'Something was approved', 'Reserved for future use'],
    ['REJECT', 'Something was rejected', 'Reserved for future use'],
    ['FILE_UPLOAD', 'A 201 file was uploaded', 'Passport, PhilHealth, Signed Contract'],
    ['SIGN_IN', 'User signed in', 'Reserved for future use'],
    ['SIGN_OUT', 'User signed out', 'Reserved for future use'],
    ['MEMBER_ADD', 'User assigned to department', 'Department membership created'],
    ['MEMBER_REMOVE', 'User removed from department', 'Department membership deleted'],
    ['ROLE_CHANGE', 'User system role changed', 'Upgrade/downgrade role'],
    ['LOGIN_FAILED', 'Failed login attempt', 'Reserved for future use'],
]

table2 = doc.add_table(rows=len(action_data) + 1, cols=3)
table2.style = 'Table Grid'
for j, h in enumerate(['Action Type', 'Meaning', 'Example Uses']):
    cell = table2.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(action_data):
    for j, val in enumerate(row):
        table2.rows[i + 1].cells[j].text = val

# ── What Gets Logged ──
doc.add_heading('WHAT GETS LOGGED', 1)

doc.add_heading('User Management', 2)
items = [
    'Creating a new user — logs email, name, role, type assigned',
    'Changing a user\'s system role (e.g., Staff → Dept Manager) — logs old and new role',
    'Changing a user\'s type (Internal Staff ↔ Virtual Assistant) — logs old and new type',
    'Enabling or disabling a user — logs the status change along with user name and email',
    'Adding a user to a department — logs the department name, position, and primary flag',
    'Removing a user from a department — logs the department and user removed',
    'Granting a temporary module role — logs the module, role, and scope',
    'Revoking a temporary module role — logs what was revoked',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Department Management', 2)
items = [
    'Creating a new department — logs department name and description',
    'Editing a department\'s name or description — logs the before and after values',
    'Toggling a department active or inactive — logs the department name and status change',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('VA Profile Management', 2)
items = [
    'Editing personal information (name, email, phone, emergency contact)',
    'Editing address (province, city, barangay, zip code, landmark)',
    'Editing employment & pay (position, level, rates, GCash, Payoneer)',
    'Editing socials (Facebook name and URL)',
    'Uploading 201 files (Passport, PhilHealth, Signed Contract) — logs the VA name, file name, and Drive folder path',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── How It Works ──
doc.add_heading('HOW IT WORKS', 1)
doc.add_paragraph(
    'The audit logging system is built on three layers:'
)

doc.add_heading('1. Automatic Capture', 2)
doc.add_paragraph(
    'Every server action that modifies data automatically records an audit entry. The system '
    'captures the state of the record BEFORE the change and AFTER the change, storing both as '
    'JSON objects. This means you can always see exactly what field was modified and what the '
    'old value was.'
)

doc.add_heading('2. Who, What, When', 2)
doc.add_paragraph(
    'Each log entry records:\n'
    '\u2022 Who performed the action (actor\'s name, email, and ID)\n'
    '\u2022 What was affected (the type of record and its ID)\n'
    '\u2022 What type of action it was (Create, Update, Delete, etc.)\n'
    '\u2022 What changed (the before and after values)\n'
    '\u2022 When it happened (date and time, to the second)\n'
    '\u2022 Additional context (department affected, related names/emails)'
)

doc.add_heading('3. Silent Operation', 2)
doc.add_paragraph(
    'Audit logging runs silently in the background. If logging fails for any reason, the '
    'intended action (e.g., saving a user\'s role) still succeeds — the system never blocks '
    'a user operation because of a logging error. Failed log attempts are printed to the '
    'server console for investigation.'
)

# ── Audit Viewer ──
doc.add_heading('AUDIT LOG VIEWER', 1)
doc.add_paragraph('Location: Admin Panel → /admin/audit (Super Admin & System Admin only)')

doc.add_heading('Features', 2)
viewer_features = [
    'Shows the 100 most recent audit events (paginated in future)',
    'Search by actor name or email — find all actions by a specific user',
    'Filter by action type — see only role changes, only file uploads, etc.',
    'Filter by entity type — see only User changes, only Department changes, etc.',
    'Clear button to reset all filters',
    'Each event shows a color-coded badge for the action type',
    'Metadata tags show contextual information (department name, file path, user email)',
    'Expandable "Changed Fields" section shows the before vs after values side by side',
    'Timestamp shows date and time in readable format',
    'Skeleton loading animation appears while data loads',
]
for f in viewer_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Reading the Viewer', 2)
doc.add_paragraph(
    'Each row in the audit viewer represents one action:\n\n'
    '  [Avatar]   John Smith   Role Changed   User   2:15 PM\n'
    '            email: john@..., name: John Smith\n'
    '  \u25B6 Changed fields\n'
    '       Before: {"systemRole": "STAFF"}\n'
    '       After:  {"systemRole": "DEPT_MANAGER"}'
)
doc.add_paragraph(
    'This tells you that John Smith\'s role was changed from Staff to Dept Manager, '
    'along with exactly when it happened and who did it.'
)

# ── Business Rules ──
doc.add_heading('BUSINESS RULES', 1)
rules = [
    'All create, update, delete, and status-change operations in the system are automatically logged.',
    'Audit logs are immutable — once written, they cannot be modified or deleted by any user.',
    'Logs are retained permanently. There is currently no automatic cleanup or archiving.',
    'Only Super Admins and System Admins can view the audit log viewer.',
    'No user action is required to enable logging — it is always on.',
    'If an audit log fails to write, the original operation still succeeds (fail-safe design).',
    'The audit log captures the actor\'s user ID, which remains valid even if the user is later deleted.',
    'Before/after values are stored as structured JSON, not as plain text — this enables future analysis and reporting.',
    'File uploads include metadata about the folder path in Google Drive for traceability.',
    'Department-scoped actions include the department ID, enabling filtering by department in future views.',
    'The audit viewer shows a maximum of 100 entries per page to maintain performance.',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

# ── System Testing ──
doc.add_heading('SYSTEM TESTING', 1)

tests = [
    ['TC-A01', 'User creation is logged', 'Create a new user via Add User form', 'Audit log shows CREATE event with email, name, role, type', 'Pass'],
    ['TC-A02', 'Role change is logged', 'Change a user\'s system role in UserCard', 'Audit log shows ROLE_CHANGE with before (old role) and after (new role)', 'Pass'],
    ['TC-A03', 'Type change is logged', 'Change a user\'s type in UserCard', 'Audit log shows UPDATE with before (old type) and after (new type)', 'Pass'],
    ['TC-A04', 'User enable/disable is logged', 'Click enable/disable button on a user', 'Audit log shows STATUS_CHANGE with user name and email', 'Pass'],
    ['TC-A05', 'Department creation is logged', 'Create a new department via inline form', 'Audit log shows CREATE event with department name', 'Pass'],
    ['TC-A06', 'Department edit is logged', 'Edit a department name via inline edit', 'Audit log shows UPDATE with before (old name) and after (new name)', 'Pass'],
    ['TC-A07', 'Department toggle is logged', 'Toggle department active/inactive', 'Audit log shows STATUS_CHANGE with department name', 'Pass'],
    ['TC-A08', 'Membership add is logged', 'Assign a user to a department', 'Audit log shows MEMBER_ADD with department name', 'Pass'],
    ['TC-A09', 'Membership remove is logged', 'Remove a user from a department', 'Audit log shows MEMBER_REMOVE with department name and user email', 'Pass'],
    ['TC-A10', 'Temp role grant is logged', 'Grant a temporary module role', 'Audit log shows CREATE with module name, role, and scope', 'Pass'],
    ['TC-A11', 'Temp role revoke is logged', 'Revoke a temporary module role', 'Audit log shows DELETE with module and role details', 'Pass'],
    ['TC-A12', 'VA profile edit is logged', 'Edit personal info / address / employment / socials', 'Audit log shows UPDATE with before/after field values', 'Pass'],
    ['TC-A13', 'File upload is logged', 'Upload a passport photo or signed contract', 'Audit log shows FILE_UPLOAD with VA name, file name, and Drive folder path', 'Pass'],
    ['TC-A14', 'Audit viewer loads with data', 'Navigate to /admin/audit after performing actions', 'Entries appear with actor name, action badge, timestamp, and metadata', 'Pass'],
    ['TC-A15', 'Filter by action type', 'Select "Role Changed" from the Action filter', 'Only ROLE_CHANGE events are shown', 'Pass'],
    ['TC-A16', 'Filter by entity type', 'Select "Department" from the Entity filter', 'Only department-related events are shown', 'Pass'],
    ['TC-A17', 'Search by actor name', 'Type an admin\'s name in the search box; press Enter', 'Only events performed by that admin are shown', 'Pass'],
    ['TC-A18', 'Clear filters', 'Apply filters; click Clear button', 'All filters reset; all events shown', 'Pass'],
    ['TC-A19', 'Changed fields expandable', 'Click "Changed fields" on an event with before/after data', 'Before section (red) and After section (green) displayed side by side', 'Pass'],
    ['TC-A20', 'Skeleton loading shown', 'Navigate to /admin/audit on a slow connection', 'Skeleton placeholder rows animate while data loads', 'Pass'],
    ['TC-A21', 'Event count displayed', 'View the audit log page', 'Header shows total number of recorded events', 'Pass'],
    ['TC-A22', 'Metadata tags visible', 'View a file upload event', 'Tags show VA name, folder path, and file name', 'Pass'],
]

table3 = doc.add_table(rows=len(tests) + 1, cols=5)
table3.style = 'Table Grid'
for j, h in enumerate(['Test ID', 'Test Scenario', 'Test Steps', 'Expected Result', 'Status']):
    cell = table3.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(tests):
    for j, val in enumerate(row):
        table3.rows[i + 1].cells[j].text = val

# ── Non-Functional ──
doc.add_heading('NON-FUNCTIONAL CONSIDERATIONS', 1)
nf = [
    'Performance: Audit write operations complete in milliseconds and do not block the main action. The viewer queries the latest 100 entries which loads efficiently even with thousands of records.',
    'Reliability: Failed audit log writes are caught silently — they never cause user-facing errors. Errors are logged to the server console for monitoring.',
    'Security: The audit viewer is restricted to Super Admin and System Admin roles via requireSuperAdmin(). All other roles are redirected to the dashboard.',
    'Scalability: The AuditLog table is indexed on (entityType, entityId, createdAt) and (actorId, createdAt) for efficient querying and filtering.',
    'Data Integrity: Before values are captured before the update and after values reflect the final state. In multi-step operations (e.g., employment form saves both VA profile and user profile), each step generates its own audit entry.',
    'Privacy: Audit logs may contain personally identifiable information (names, emails). Access to the viewer is strictly role-restricted.',
]
for item in nf:
    doc.add_paragraph(item, style='List Bullet')

# ── Files Reference ──
doc.add_heading('FILES REFERENCE', 1)
files_ref = [
    ('lib/audit.ts', 'Core audit logging utility. Provides logAudit() and logAuditWrite() functions.'),
    ('prisma/schema.prisma', 'AuditLog model and AuditAction enum definition.'),
    ('app/(dashboard)/admin/audit/page.tsx', 'Audit log viewer page with filters, color-coded badges, and expandable diffs.'),
    ('app/(dashboard)/admin/audit/loading.tsx', 'Skeleton loading state for the audit viewer.'),
    ('app/(dashboard)/admin/users/actions.ts', 'All user and department actions wired with audit logging.'),
    ('app/(dashboard)/vas/actions.ts', 'VA profile and employment actions wired with audit logging.'),
    ('app/api/upload/route.ts', 'File upload endpoint wired with FILE_UPLOAD audit logging.'),
    ('components/vas/VAProfileEditor.tsx', 'Passes current user ID to upload API for actor tracking.'),
    ('scripts/add_audit_actions.js', 'Database migration script for adding new audit action enum values.'),
]
for path, desc in files_ref:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.bold = True
    p.add_run('\n' + desc)

p = doc.add_paragraph('\n\u2014 End of Document \u2014')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('Audit-Logs-Module-Documentation.docx')
print('Saved Audit-Logs-Module-Documentation.docx')
