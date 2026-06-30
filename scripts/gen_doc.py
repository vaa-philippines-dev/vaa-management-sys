from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

t = doc.add_heading('VAA Management System \u2014 Prototype Update', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('June 29, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

doc.add_heading('Executive Summary', 1)
doc.add_paragraph(
    'This document outlines the new features, database additions, and role-based '
    'access control system added to the VAA Management System prototype.'
)

doc.add_heading('1. Role-Based Access Control', 1)
table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
for j, h in enumerate(['Role', 'Create', 'Read', 'Update', 'Delete']):
    table.rows[0].cells[j].text = h
    for p in table.rows[0].cells[j].paragraphs:
        for r in p.runs: r.bold = True
for i, row in enumerate([
    ['HR', 'Yes', 'Yes', 'Yes', 'No'],
    ['SERVICE / CS', 'No', 'Yes', 'No', 'No'],
]):
    for j, v in enumerate(row):
        table.rows[i+1].cells[j].text = v

doc.add_heading('2. Database Additions \u2014 UserProfile', 2)
up_fields = [
    ('gender', 'Text', 'Manual'),
    ('whatsapp_number', 'Text', 'Manual'),
    ('non_celebrant', 'Boolean', 'Manual'),
    ('barangay', 'Text', 'Manual'),
    ('city_municipality', 'Text', 'Manual'),
    ('province', 'Text', 'Manual'),
    ('zip_code', 'Text', 'Manual'),
    ('landmark', 'Text', 'Manual'),
    ('passport_number', 'Text', 'Manual'),
    ('passport_photo', 'Text', 'Manual'),
    ('philhealth_number', 'Text', 'Manual'),
    ('philhealth_photo', 'Text', 'Manual'),
]
t2 = doc.add_table(rows=len(up_fields)+1, cols=3)
t2.style = 'Table Grid'
for j, h in enumerate(['Field', 'Type', 'Access']):
    t2.rows[0].cells[j].text = h
    for p in t2.rows[0].cells[j].paragraphs:
        for r in p.runs: r.bold = True
for i, (f, ty, acc) in enumerate(up_fields):
    t2.rows[i+1].cells[0].text = f
    t2.rows[i+1].cells[1].text = ty
    t2.rows[i+1].cells[2].text = acc

doc.add_heading('VAProfile \u2014 New Fields', 2)
vp_fields = [
    ('base_rate', 'Decimal', 'PHP hourly rate'),
    ('vaa_position', 'Text', 'VAA position title'),
    ('level', 'Text', 'Skill level'),
    ('preferred_work_hours', 'Decimal', 'Weekly hours'),
    ('available_schedule', 'Text', 'Schedule'),
    ('hybrid', 'Boolean', 'Multi-department'),
    ('contract_link', 'Text', 'Contract file'),
    ('folder_201_link', 'Text', '201 folder'),
    ('file_201_link', 'Text', '201 file'),
    ('va_client_file_link', 'Text', 'VA-client file'),
    ('health_check_file_link', 'Text', 'Health check'),
    ('va_profile_link', 'Text', 'VA profile doc'),
    ('payout_summary_link', 'Text', 'Payout summary'),
    ('dept_201_folder_link', 'Text', 'Dept 201 folder'),
]
t3 = doc.add_table(rows=len(vp_fields)+1, cols=3)
t3.style = 'Table Grid'
for j, h in enumerate(['Field', 'Type', 'Description']):
    t3.rows[0].cells[j].text = h
    for p in t3.rows[0].cells[j].paragraphs:
        for r in p.runs: r.bold = True
for i, (f, ty, desc) in enumerate(vp_fields):
    t3.rows[i+1].cells[0].text = f
    t3.rows[i+1].cells[1].text = ty
    t3.rows[i+1].cells[2].text = desc

doc.add_heading('DocumentType Enum \u2014 New Values', 2)
for v in ['FOLDER_201', 'FILE_201', 'VA_CLIENT_FILE', 'HEALTH_CHECK', 'VA_PROFILE', 'PAYOUT_SUMMARY', 'DEPT_201_FOLDER']:
    doc.add_paragraph(v, style='List Bullet')

doc.add_heading('3. New Features', 1)

doc.add_heading('VA Master List (/vas)', 2)
for item in [
    'Stats: Total VAs, Active, Available, Assignments count',
    'Table: Name, Email, Dept/Position, VAA Position, Contract, Availability, Rate, Assignments',
    'HR view shows personal data; SERVICE/CS view is read-only basics',
    'Click rows to open full VA profile',
    'Assignments as clickable client badge chips',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('VA Profile Detail (/vas/[id])', 2)
for item in [
    'Personal Info: name, gender, birthday, WhatsApp, GCASH',
    'Address: structured (barangay, city, province, zip, landmark)',
    'Employment: contract type, status, dates, VAA position, level, rates',
    'Emergency Contact, Socials, Identification (passport, PhilHealth)',
    'Documents section with all Google Drive links and uploaded files',
    'Skills badges and Assignment cards with status',
    'HR can click Edit on any section to inline-edit and Save',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Admin Panel (/admin)', 2)
for item in [
    'Overview stats, department grid, role distribution chart, recent users',
    'Quick actions: Manage Users, Department Views, VA Workforce, Reports',
    'Click any department card for filtered dashboard view',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('UX Improvements', 1)
for item in [
    'Skeleton loading screens during navigation',
    'Click animations: scale(0.97) on press for all buttons/links',
    'Card hover effects: lift on hover, press down on click',
    'Smooth CSS transitions on all interactive elements',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4. Bug Fixes', 1)
for item in [
    'Cookie modification error in Next.js 16 (proxy.ts session refresh)',
    'Database connection pool exhaustion (transaction-mode pgbouncer)',
    'OAuth redirect URL on production (window.location.origin)',
    'Google account chooser on every sign-in (prompt=select_account)',
    'PostgreSQL enum types creation for all Prisma enums',
    'createdAt/updatedAt column migration (camelCase to snake_case)',
    'VA profile user_id links for orphan profiles',
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph('\n\u2014 End of Document \u2014')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('VAA-Prototype-Documentation.docx')
print('Saved VAA-Prototype-Documentation.docx')
