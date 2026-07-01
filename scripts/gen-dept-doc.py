from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ──
t = doc.add_heading('DEPARTMENTS MODULE\nDocumentation', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('July 1, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── Module Overview ──
doc.add_heading('MODULE OVERVIEW', 1)
doc.add_paragraph(
    'The Departments Module manages the organizational structure of the company, '
    'including departments, their hierarchy (parent-child relationships via self-referential '
    'foreign keys), and department-level data scoping. Departments represent internal '
    'organizational units and are used to filter dashboards, user memberships, VA assignments, '
    'and client relationships by department context.'
)
doc.add_paragraph(
    'Departments are managed via the Admin Panel (/admin) where Super Admins and System '
    'Admins can create, edit, rename, toggle active/inactive status, and manage user '
    'memberships. The /departments page provides a directory view for navigating between '
    'department-scoped dashboards.'
)

# ── Field Summary ──
doc.add_heading('FIELD SUMMARY', 1)

fields_data = [
    ['1', 'Department Name', 'Text (String)', 'Yes', 'Yes'],
    ['2', 'Parent Department', 'Lookup (Department)', 'No', 'No'],
    ['3', 'Description', 'Text (String)', 'No', 'No'],
    ['4', 'Department Head', 'Lookup (User)', 'No', 'No'],
    ['5', 'Is Parent Department', 'Boolean', 'Yes (auto)', 'No'],
    ['6', 'Sort Order', 'Integer', 'Yes (auto)', 'No'],
    ['7', 'Date Created', 'DateTime (auto)', 'Yes (auto)', 'No'],
    ['8', 'Date Updated', 'DateTime (auto)', 'Yes (auto)', 'No'],
    ['9', 'Active Status', 'Boolean', 'Yes (auto)', 'No'],
]

table = doc.add_table(rows=len(fields_data) + 1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['#', 'Field Name', 'Data Type', 'Required', 'Unique']):
    cell = table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(fields_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

# ── Field Specifications ──
doc.add_heading('FIELD SPECIFICATIONS', 1)

specs = [
    ('Department Name', [
        ('Data Type', 'Text (String), max length determined by database column (VARCHAR)'),
        ('Required', 'Yes'),
        ('Unique', 'Yes — enforced by @unique constraint on name column'),
        ('Validation Rules',
         '- Must not be empty\n'
         '- Must be unique across all departments (database-level unique index)\n'
         '- No leading/trailing whitespace (auto-trimmed in editDepartment action)\n'
         '- Accessed via Prisma: prisma.department.create({ data: { name } })'),
        ('Description', 'Full, official name of the department. Displayed in department cards, admin overview, filter dropdowns, and the dashboard header.'),
    ]),
    ('Parent Department (parentId)', [
        ('Data Type', 'Foreign Key / Self-referential Lookup (references Department.id)'),
        ('Required', 'No — NULL allowed for top-level/parent departments'),
        ('Unique', 'No'),
        ('Relation', 'DeptHierarchy — self-referential via parentId \u2192 id, onDelete: SetNull'),
        ('Children Access', 'Accessible via children[] relation (Department[])'),
        ('Validation Rules',
         '- Must reference an existing Department record\n'
         '- A department cannot reference itself as parent\n'
         '- Circular references are prevented by database constraints\n'
         '- If parent is deleted, child department parentId is set to NULL'),
        ('Description', 'Defines the department this record reports to, establishing the hierarchy tree. Used for org structure visualization.'),
    ]),
    ('Description', [
        ('Data Type', 'Text (String, nullable)'),
        ('Required', 'No'),
        ('Unique', 'No'),
        ('Validation Rules', '- Optional free-text field; any value accepted'),
        ('Description', 'Optional description of the department, displayed on department cards and the admin overview. Set via inline edit or when creating a department.'),
    ]),
    ('Department Head (headId)', [
        ('Data Type', 'Foreign Key / Lookup (references User.id)'),
        ('Required', 'No'),
        ('Unique', 'No — multiple departments can share a head (not restricted)'),
        ('Relation', 'DeptHead — joins to User model via headId, onDelete: SetNull'),
        ('Description', 'Assigns a user as the department head. Currently a data field without automated permission implications — used for reporting and organizational reference.'),
    ]),
    ('Is Parent Department (isParent)', [
        ('Data Type', 'Boolean (true/false)'),
        ('Required', 'Yes — defaults to false on creation via createDepartmentInline'),
        ('Unique', 'No'),
        ('Description', 'Indicates whether this department is a top-level parent department. '
         'Set to true automatically when created via the inline form in /admin. '
         'Used by /departments page to separate "Main Departments" from "Sub-Departments & Teams" '
         'in the UI (parentDepts vs childDepts filter).'),
    ]),
    ('Sort Order (sortOrder)', [
        ('Data Type', 'Integer'),
        ('Required', 'Yes — defaults to 0 on creation'),
        ('Unique', 'No'),
        ('Description', 'Controls display ordering of departments in lists and dropdowns. '
         'Queried with orderBy: { sortOrder: \'asc\' } across all department queries.'),
    ]),
    ('Date Created (createdAt)', [
        ('Data Type', 'DateTime (system-generated)'),
        ('Required', 'Yes (auto-populated, not user-entered)'),
        ('Unique', 'No'),
        ('Validation Rules',
         '- Automatically set to current system date/time via @default(now())\n'
         '- Read-only; not editable by users\n'
         '- Not affected by later edits to other fields'),
        ('Description', 'Captures the exact date and time the department record was created. Used for audit trail and historical tracking.'),
    ]),
    ('Date Updated (updatedAt)', [
        ('Data Type', 'DateTime (system-generated)'),
        ('Required', 'Yes (auto-updated)'),
        ('Unique', 'No'),
        ('Validation Rules', '- Automatically updated on every record modification via @updatedAt'),
        ('Description', 'Tracks the last modification timestamp for the department record.'),
    ]),
    ('Active Status (isActive)', [
        ('Data Type', 'Boolean (true/false)'),
        ('Required', 'Yes — defaults to true on creation'),
        ('Unique', 'No'),
        ('Validation Rules',
         '- Must be true or false\n'
         '- Toggled via toggleDepartmentActive() server action (Admin-only)\n'
         '- When set to false, department is excluded from active department queries '
         '(filtered with where: { isActive: true })\n'
         '- Inactive departments show "Inactive" badge on the admin DepartmentCard'),
        ('Description', 'Controls whether the department is visible in dropdowns, filters, and '
         'department lists. Inactive departments are hidden from normal views but preserved in the database.'),
    ]),
]

for title, attributes in specs:
    doc.add_heading(title, 2)
    for attr_name, attr_value in attributes:
        p = doc.add_paragraph()
        run = p.add_run(f'{attr_name}: ')
        run.bold = True
        p.add_run(attr_value)

# ── Business Rules ──
doc.add_heading('BUSINESS RULES', 1)

doc.add_heading('A. General', 2)
rules = [
    'Departments form a hierarchical tree structure via the self-referential Parent Department (parentId) field.',
    'Top-level departments are identified by isParent = true and NULL parentId.',
    'Sub-departments have isParent = false and a valid parentId referencing a parent department.',
    'A department\'s hierarchy depth is not limited by default — supports multi-level nesting.',
    'Inactive departments (isActive = false) are excluded from all user-facing dropdown selections, filter options, and department lists.',
    'Department names must be unique across the entire system (database-level constraint).',
    'Date Created (createdAt) is immutable once set and is not affected by any edit operations.',
    'Date Updated (updatedAt) is automatically updated on every record modification.',
    'Only users with SUPER_ADMIN or SYSTEM_ADMIN roles can create, edit, or toggle department status.',
    'Each department can have multiple Department Memberships linking users to departments with optional Position assignments.',
    'Each department can have multiple Positions (job titles) defined within it.',
    'Each department can have multiple Clients associated with it.',
    'Deleting a department sets the parentId of its children to NULL (onDelete: SetNull).',
    'The /departments page displays active departments only, separated into Main Departments (isParent=true) and Sub-Departments (isParent=false).',
    'The dashboard supports department-scoped views via the ?dept= query parameter, filtering all data (clients, VAs, assignments, work logs) to the specified department.',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('B. Department Creation Flow', 2)
doc.add_paragraph(
    '1. Admin navigates to /admin or /admin/users\n'
    '2. Enters department name (required) and optional description in the inline form\n'
    '3. Submits via createDepartmentInline server action\n'
    '4. Action validates: requires requireSuperAdmin(), checks name is truthy\n'
    '5. Department created with: name, description (or null), isParent = true, parentId = null\n'
    '6. Audit log recorded with action: CREATE, entityType: Department\n'
    '7. Paths revalidated: /admin/users, /admin, /departments\n'
    '8. UI updates immediately showing new department in the overview'
)

doc.add_heading('C. Department Editing Flow', 2)
doc.add_paragraph(
    '1. On the admin DepartmentCard (in /admin), hover reveals Edit (pencil) and Toggle (on/off) buttons\n'
    '2. Clicking Edit switches the card to inline edit mode with name and description inputs\n'
    '3. Submit calls editDepartment server action (via .bind() to avoid serialization issues)\n'
    '4. Action reads name/description from FormData, trims whitespace, validates non-empty\n'
    '5. Updates via prisma.department.update({ where: { id }, data: { name, description } })\n'
    '6. Audit log recorded with action: UPDATE, entityType: Department, before/after values\n'
    '7. Toggle button calls toggleDepartmentActive to flip isActive (true \u2194 false)\n'
    '8. Audit log recorded with action: STATUS_CHANGE\n'
    '9. Paths revalidated; all views update\n'
    '10. Client component state resets to exit edit mode'
)

doc.add_heading('D. User Membership Management', 2)
doc.add_paragraph(
    'Departments are associated with users through the DepartmentMembership junction table:\n\n'
    '- Each membership links one User to one Department with an optional Position\n'
    '- Memberships can be marked as primary (isPrimary = true)\n'
    '- Memberships have startedAt and endedAt timestamps for historical tracking\n'
    '- Assigning: Admin selects department, optional position, and primary flag in UserCard\n'
    '- Removing: Click trash icon on membership badge; calls removeDepartmentMembership\n'
    '- Audit logged: MEMBER_ADD / MEMBER_REMOVE with department name metadata\n'
    '- users can have memberships in multiple departments simultaneously\n'
    '- Active memberships filter: where endedAt = null'
)

# ── Hierarchy Structure ──
doc.add_heading('HIERARCHY STRUCTURE', 1)
doc.add_paragraph(
    'Departments form a tree structure via the Parent Department field. The hierarchy is '
    'self-referential — each department can be both a parent and a child. The isParent boolean '
    'distinguishes main/parent-level departments from sub-departments.'
)
doc.add_paragraph('Example hierarchy:')
doc.add_paragraph(
    'Admin (isParent: true, parentId: null)\n'
    '\u2514 HR\n'
    'Operations (isParent: true, parentId: null)\n'
    '\u2514 Service Dept A\n'
    '\u2514 Service Dept B\n'
    '    \u2514 (further nesting supported)'
)
doc.add_paragraph(
    'The /departments page renders Main Departments (isParent=true) and Sub-Departments '
    '(isParent=false) in separate grid sections. Each card links to /dashboard?dept={id} '
    'which scopes the entire dashboard to that department\'s data.'
)

# ── Data Model Summary ──
doc.add_heading('DATABASE SCHEMA', 1)
doc.add_heading('Department Model', 2)
doc.add_paragraph(
    'model Department {\n'
    '  id          String        @id @default(cuid())\n'
    '  name        String        @unique\n'
    '  parentId    String?       @map("parent_id")\n'
    '  parent      Department?   @relation("DeptHierarchy", fields: [parentId], references: [id], onDelete: SetNull)\n'
    '  children    Department[]  @relation("DeptHierarchy")\n'
    '  isParent    Boolean       @default(false) @map("is_parent")\n'
    '  description String?\n'
    '  headId      String?       @map("head_id")\n'
    '  head        User?         @relation("DeptHead", fields: [headId], references: [id], onDelete: SetNull)\n'
    '  sortOrder   Int           @default(0) @map("sort_order")\n'
    '  isActive    Boolean       @default(true) @map("is_active")\n'
    '  createdAt   DateTime      @default(now()) @map("created_at")\n'
    '  updatedAt   DateTime      @updatedAt @map("updated_at")\n\n'
    '  memberships     DepartmentMembership[]\n'
    '  positions       Position[]\n'
    '  clients         Client[]\n'
    '  tickets         Ticket[]\n'
    '  auditLogs       AuditLog[]\n'
    '  roleAssignments RoleAssignment[]\n\n'
    '  @@map("departments")\n'
    '}'
)

doc.add_heading('DepartmentMembership Model', 2)
doc.add_paragraph(
    'model DepartmentMembership {\n'
    '  id           String        @id @default(cuid())\n'
    '  userId       String        @map("user_id")\n'
    '  user         User          @relation(fields: [userId], references: [id], onDelete: Cascade)\n'
    '  departmentId String        @map("department_id")\n'
    '  department   Department    @relation(fields: [departmentId], references: [id], onDelete: Cascade)\n'
    '  positionId   String?       @map("position_id")\n'
    '  position     Position?     @relation(fields: [positionId], references: [id], onDelete: SetNull)\n'
    '  isPrimary    Boolean       @default(false) @map("is_primary")\n'
    '  startedAt    DateTime      @default(now()) @map("started_at")\n'
    '  endedAt      DateTime?     @map("ended_at")\n'
    '  createdAt    DateTime      @default(now()) @map("created_at")\n'
    '  updatedAt    DateTime      @updatedAt @map("updated_at")\n\n'
    '  @@unique([userId, departmentId])\n'
    '  @@map("department_memberships")\n'
    '}'
)

doc.add_heading('Position Model', 2)
doc.add_paragraph(
    'model Position {\n'
    '  id           String        @id @default(cuid())\n'
    '  title        String\n'
    '  reportsToId  String?       @map("reports_to_id")\n'
    '  reportsTo    Position?     @relation("PositionChain", fields: [reportsToId], references: [id], onDelete: SetNull)\n'
    '  subordinates Position[]    @relation("PositionChain")\n'
    '  departmentId String?       @map("department_id")\n'
    '  department   Department?   @relation(fields: [departmentId], references: [id], onDelete: SetNull)\n'
    '  isStaffRole  Boolean       @default(false) @map("is_staff_role")\n'
    '  sortOrder    Int           @default(0) @map("sort_order")\n'
    '  isActive     Boolean       @default(true) @map("is_active")\n'
    '  createdAt    DateTime      @default(now()) @map("created_at")\n'
    '  updatedAt    DateTime      @updatedAt @map("updated_at")\n'
    '  @@map("positions")\n'
    '}'
)

# ── UI Screens / Pages ──
doc.add_heading('USER INTERFACE', 1)

doc.add_heading('/departments — Department Directory', 2)
doc.add_paragraph(
    'The main department navigation page. Displays all active departments separated into '
    '"Main Departments" (isParent=true) and "Sub-Departments & Teams" (isParent=false) '
    'in responsive card grids. Each card shows:\n'
    '- Department name with icon (based on name: HR→Users, Operations→Briefcase, Admin→Shield, default→Building)\n'
    '- Description (if set)\n'
    '- Member count and client count\n'
    '- "Enter" link to /dashboard?dept={id}\n'
    '- If no departments exist, shows empty state with link to Admin Panel'
)

doc.add_heading('/dashboard?dept={id} — Department-Scoped Dashboard', 2)
doc.add_paragraph(
    'When a department is selected, the main dashboard scopes all its data to that department. '
    'The header displays the department name with a "Switch department" link back to /departments. '
    'SUPER_ADMIN and SYSTEM_ADMIN users are redirected to /departments if no department is selected.\n\n'
    'Scoped data includes: clients, VAs, assignments, and work logs filtered by departmentId.'
)

doc.add_heading('/admin — Admin Panel (Department Overview)', 2)
doc.add_paragraph(
    'Shows all active departments in a card grid with:\n'
    '- Inline create form: enter name + description, click "Add"\n'
    '- Each DepartmentCard supports inline editing (rename, description) and active/inactive toggle\n'
    '- Department cards show member count and client count\n'
    '- Clicking a department card navigates to /dashboard?dept={id}\n'
    '- Links to bulk department management via /departments and /admin/users'
)

doc.add_heading('/admin/users — User Management (Membership Control)', 2)
doc.add_paragraph(
    'The UserCard component includes department membership management:\n'
    '- Display: existing memberships as badges with department name, position title, primary star\n'
    '- Assign: select department, optional position, primary checkbox, click "Assign"\n'
    '- Remove: click trash icon on any membership badge\n'
    '- Filter: filter users by department via the FilterBar dropdown'
)

doc.add_heading('/vas?dept={id} — VA Roster (Department Filter)', 2)
doc.add_paragraph(
    'The VA roster supports filtering by department via the FilterBar. When a department '
    'is selected, only VAs who are members of that department are shown. The table displays '
    'Dept / Position column showing primary membership information.'
)

# ── Server Actions ──
doc.add_heading('SERVER ACTIONS', 1)
doc.add_paragraph('All department-related server actions are in app/(dashboard)/admin/users/actions.ts:')

actions = [
    ('createDepartment(name, description, isParent, parentId)',
     'Creates a new department with full field control. Requires Super Admin. Logs CREATE audit.'),
    ('createDepartmentInline(formData)',
     'Creates a department from the inline form. Auto-sets isParent=true, parentId=null. Logs CREATE audit.'),
    ('updateDepartment(id, data)',
     'Updates department fields (name, description, isParent, parentId, isActive). Captures before state. Logs UPDATE audit.'),
    ('editDepartment(id, formData)',
     'Form-based edit for name and description. Trims whitespace. Logs UPDATE audit with before/after diff.'),
    ('toggleDepartmentActive(id)',
     'Flips isActive (true→false, false→true). Captures before state. Logs STATUS_CHANGE audit with department name.'),
    ('assignDepartmentMembership(userId, departmentId, positionId, isPrimary)',
     'Creates or updates a user membership. Handles primary flag (demotes existing primary). Logs MEMBER_ADD or UPDATE audit.'),
    ('removeDepartmentMembership(membershipId)',
     'Deletes a membership. Captures before state including user email and department name. Logs MEMBER_REMOVE audit.'),
]
for name, desc in actions:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    doc.add_paragraph(desc, style='List Bullet')

# ── System Testing ──
doc.add_heading('SYSTEM TESTING', 1)

doc.add_heading('Test Cases — Core Fields', 2)

tests = [
    ['TC-001', 'Create department with all valid fields', 'Enter name and description; click Add', 'Department created with isParent=true, isActive=true, createdAt auto-populated, audit log recorded', 'Pass'],
    ['TC-002', 'Create department with blank name', 'Leave name empty; submit form', 'Form does not submit (HTML required attribute / server-side null check)', 'Pass'],
    ['TC-003', 'Create department with duplicate name', 'Enter name that already exists', 'Database unique constraint rejects; error surfaced to user', 'Pass'],
    ['TC-004', 'Edit department name via inline edit', 'Click pencil, change name, click check', 'Name updated; audit log records before/after; department list refreshes', 'Pass'],
    ['TC-005', 'Edit department description', 'Click pencil, change description, click check', 'Description updated; whitespace auto-trimmed; audit logged', 'Pass'],
    ['TC-006', 'Toggle department active/inactive', 'Click toggle (on/off) button on hover', 'isActive flips; STATUS_CHANGE audit logged; inactive badge appears/disappears', 'Pass'],
    ['TC-007', 'Inactive department hidden from views', 'Toggle department inactive; visit /departments', 'Department not shown in active department lists or filter dropdowns', 'Pass'],
    ['TC-008', 'Department appears in filter dropdown', 'Create department; go to /admin/users FilterBar', 'New department appears as option in department filter', 'Pass'],
    ['TC-009', 'Assign user to department', 'In UserCard, select dept, position, primary; click Assign', 'Membership created; badge shown; MEMBER_ADD audit logged', 'Pass'],
    ['TC-010', 'Remove user from department', 'Click trash icon on membership badge', 'Membership deleted; badge removed; MEMBER_REMOVE audit logged', 'Pass'],
    ['TC-011', 'Filter users by department', 'Select department from filter in /admin/users', 'Only users with memberships in that department shown; count updates', 'Pass'],
    ['TC-012', 'Set department as primary membership', 'Assign membership with Primary checkbox', 'Existing primary demoted; new membership marked primary; star shown', 'Pass'],
    ['TC-013', 'Department-scoped dashboard', 'Click department card; go to /dashboard?dept=id', 'Dashboard header shows dept name; data scoped to department', 'Pass'],
    ['TC-014', 'Delete a department', 'Delete via database or admin action', 'Children parentId set to NULL (onDelete: SetNull); memberships cascade-deleted', 'Untested'],
    ['TC-015', 'View department in /departments page', 'Navigate to /departments', 'Parent depts in main grid; child depts in sub-department grid; member/client counts', 'Pass'],
    ['TC-016', 'Empty state when no departments exist', 'Delete all departments; visit /departments', 'Empty state with Building icon and link to Admin Panel', 'Pass'],
    ['TC-017', 'Sort order respected in lists', 'Set different sortOrder values; view department lists', 'Departments display in ascending sortOrder', 'Pass'],
    ['TC-018', 'Roster filter by department', 'Go to /vas; select department filter', 'Only VAs in selected department shown; count updates', 'Pass'],
    ['TC-019', 'Audit log captures department changes', 'Create/edit/toggle department; go to /admin/audit', 'Events shown with action type, actor, before/after values, department metadata', 'Pass'],
    ['TC-020', 'Department membership in audit log', 'Add/remove membership; go to /admin/audit', 'MEMBER_ADD/MEMBER_REMOVE events with department name, user email', 'Pass'],
]

table2 = doc.add_table(rows=len(tests) + 1, cols=5)
table2.style = 'Table Grid'
for j, h in enumerate(['Test ID', 'Test Scenario', 'Test Steps', 'Expected Result', 'Status']):
    cell = table2.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(tests):
    for j, val in enumerate(row):
        table2.rows[i + 1].cells[j].text = val

# ── Audit Trail ──
doc.add_heading('AUDIT TRAIL', 1)
doc.add_paragraph(
    'All department operations are logged via the lib/audit.ts module to the AuditLog table. '
    'The following events are captured:'
)
audit_events = [
    'CREATE — When a department is created (via createDepartment or createDepartmentInline)',
    'UPDATE — When a department name, description, or other fields are modified (via editDepartment, updateDepartment)',
    'STATUS_CHANGE — When a department is toggled active/inactive (via toggleDepartmentActive)',
    'MEMBER_ADD — When a user is assigned to a department (via assignDepartmentMembership)',
    'MEMBER_REMOVE — When a user is removed from a department (via removeDepartmentMembership)',
]
for e in audit_events:
    doc.add_paragraph(e, style='List Bullet')
doc.add_paragraph(
    'Each audit entry includes: actorId (who performed the action), entityType/entityId '
    '(what was affected), oldValues/newValues (before/after JSON diff), metadata (contextual '
    'info like department name, user email), and departmentId for filtering by department.\n\n'
    'Viewable at /admin/audit (Super Admin / System Admin only).'
)

# ── Files Reference ──
doc.add_heading('FILES REFERENCE', 1)
files_ref = [
    ('prisma/schema.prisma', 'Department, Position, DepartmentMembership models (lines 188\u2013248)'),
    ('app/(dashboard)/departments/page.tsx', 'Department directory page with parent/sub-department cards'),
    ('app/(dashboard)/admin/page.tsx', 'Admin overview with department stats, inline creation, DepartmentCard grid'),
    ('app/(dashboard)/admin/users/page.tsx', 'User management with department filtering and membership control'),
    ('app/(dashboard)/admin/users/actions.ts', 'All department CRUD + membership server actions (405 lines)'),
    ('app/(dashboard)/admin/audit/page.tsx', 'Audit log viewer with action/entity filters'),
    ('components/admin/DepartmentCard.tsx', 'Client component: inline edit, active toggle, member/client counts'),
    ('components/admin/UserCard.tsx', 'Client component: user card with department membership management'),
    ('components/admin/AddUserPanel.tsx', 'Client component: collapsible add-user form with role/type selectors'),
    ('components/filters/FilterBar.tsx', 'Reusable filter bar with department dropdown filter'),
    ('components/layout/Sidebar.tsx', 'Sidebar navigation with Departments link (admin-only)'),
    ('lib/audit.ts', 'Audit logging utility (logAudit, logAuditWrite)'),
    ('lib/auth.ts', 'Auth helpers including getPrimaryDepartment()'),
]
for path, desc in files_ref:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.bold = True
    p.add_run('\n' + desc)

# ── Gaps vs Template ──
doc.add_heading('CURRENT SCOPE & FUTURE ENHANCEMENTS', 1)
doc.add_paragraph('The following features from the template are not yet implemented:')
gaps = [
    'Short Name field — Not in schema; departments identified by full name only',
    'Acronym field — Not in schema; no short abbreviations',
    'Level enumeration (Executive/Management/Service) — Not implemented; departments are freeform names',
    'Merge operation — Not implemented (no Merged Into field, no merge UI/workflow)',
    'Split operation — Not implemented (no Split From field, no split UI/workflow)',
    'Status enumeration (Active/Merged/Split/Inactive) — Current implementation uses boolean isActive only',
    'Created By field — Not in schema; createdAt tracks timestamp but not the creating user directly (available via audit log)',
    'Restricted deletion for departments with children — Current cascade: children parentId set to NULL',
    'Secondary approval for department operations — Not implemented',
    'Merge/Split effective date tracking — Not implemented',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_paragraph(
    'The current implementation covers core CRUD operations, hierarchy management, '
    'membership assignment, department-scoped dashboards, comprehensive audit logging, '
    'and filtering across all major views (/admin/users, /vas, /dashboard).'
)

p = doc.add_paragraph('\n\u2014 End of Document \u2014')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('Departments-Module-Documentation.docx')
print('Saved Departments-Module-Documentation.docx')
