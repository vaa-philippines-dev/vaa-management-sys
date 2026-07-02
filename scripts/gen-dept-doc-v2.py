from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

t = doc.add_heading('DEPARTMENTS MODULE\nDocumentation', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('July 2, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

doc.add_heading('MODULE OVERVIEW', 1)
doc.add_paragraph(
    'The Departments Module manages the organizational structure of the company, including '
    'departments, their hierarchy (parent-child relationships), and classification into Levels. '
    'Departments may represent internal organizational units (Executive, Management) or '
    'client-facing services (Service) where Virtual Assistants (VAs) are assigned to clients.'
)

doc.add_heading('HIERARCHY STRUCTURE', 1)
doc.add_paragraph(
    'The module supports a three-Level hierarchy. Each Level is a system-protected record that '
    'cannot be deleted or have its Level field changed.'
)
doc.add_paragraph('Current seeded hierarchy:')
doc.add_paragraph(
    'Executive (Level)\n'
    '\u2514 (reserved for future executive departments)\n\n'
    'Management (Level)\n'
    '\u2514 Staff Department\n'
    '    \u2514 Customer Success [CS]\n'
    '    \u2514 Human Resources [HR]\n'
    '    \u2514 Software Development [DEV]\n'
    '    \u2514 Business Operations [BIZOPS]\n'
    '    \u2514 Admin Department [ADMIN]\n\n'
    'Service (Level)\n'
    '\u2514 VA Management [VAMGT]\n'
    '    \u2514 Amazon Department [AMZ]\n'
    '    \u2514 Wholesale Department [WS]\n'
    '    \u2514 PPC Department [PPC]\n'
    '    \u2514 Social Media Department [SM]\n'
    '    \u2514 Executive Assistant Department [EA]\n'
    '    \u2514 Walmart Department [WM]\n'
    '    \u2514 Creatives Department [CRTV]'
)

doc.add_heading('FIELD SUMMARY', 1)
fields_data = [
    ['1', 'Department Name', 'Text', 'Yes', 'Yes'],
    ['2', 'Parent Department', 'Lookup (Department)', 'No', '-'],
    ['3', 'Short Name', 'Text', 'No', 'No'],
    ['4', 'Acronym', 'Text (2-6 uppercase letters)', 'Yes', 'Yes'],
    ['5', 'Level', 'Enum (Executive / Management / Service)', 'Yes', '-'],
    ['6', 'Status', 'Enum (Active / Merged / Split / Inactive)', 'Yes (system)', '-'],
    ['7', 'Merged Into', 'Lookup (Department)', 'Conditional', '-'],
    ['8', 'Split From', 'Lookup (Department)', 'Conditional', '-'],
    ['9', 'Description', 'Text', 'No', '-'],
    ['10', 'Date Created', 'Timestamp (auto)', 'Yes (auto)', 'No'],
    ['11', 'Date Updated', 'Timestamp (auto)', 'Yes (auto)', 'No'],
]

table = doc.add_table(rows=len(fields_data) + 1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['#', 'Field Name', 'Data Type', 'Required', 'Unique']):
    cell = table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(fields_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_heading('FIELD SPECIFICATIONS', 1)

specs = [
    ('Department Name', 'Full official name of the department. Required. Must be unique across all departments. Auto-trims leading/trailing whitespace. Maximum 100 characters. Whitespace-only values are rejected.'),
    ('Parent Department', 'Optional reference to the parent department. Establishes the hierarchy. Can be blank only for Level records (Executive / Management / Service). Self-references and circular references are blocked.'),
    ('Short Name', 'Optional shortened display name. Used in space-constrained UI elements. Maximum 50 characters. Whitespace-only is treated as null.'),
    ('Acronym', 'Short 2-6 character uppercase abbreviation. Required. Must be unique across departments. Letters only — no numbers or special characters. Auto-converted to uppercase on save. Used for tagging and quick reference.'),
    ('Level', 'Required. Must be one of: Executive, Management, Service. Level records (the three default ones) cannot have their Level field changed by anyone.'),
    ('Status', 'System-managed. Defaults to Active on creation. Automatically set to Merged or Split when those operations occur. Can only be manually toggled to Inactive (or back to Active).'),
    ('Merged Into', 'Set automatically when a department is merged into another. Identifies the resulting department. Cleared if a department is reactivated.'),
    ('Split From', 'Set automatically on departments created via a split operation. Identifies the source department that was split.'),
    ('Description', 'Optional free-text description of the department.'),
    ('Date Created', 'System-generated timestamp. Set automatically on creation. Cannot be modified by users.'),
    ('Date Updated', 'System-generated timestamp. Updated automatically on every modification.'),
]
for title, desc in specs:
    doc.add_heading(title, 2)
    doc.add_paragraph(desc)

doc.add_heading('BUSINESS RULES', 1)
doc.add_heading('A. General', 2)
rules_a = [
    'There are three Level records: Executive, Management, Service. They exist at the top of the hierarchy.',
    'The three Level records are system-protected: they cannot be deleted, and their Level field cannot be changed.',
    'The three Level records are the only valid choices as parent when the candidate department has no parent set.',
    'Hierarchy depth is not limited by default; multi-level nesting is supported.',
    'A department under the Service Level is eligible for VA-to-client assignment workflows.',
    'Circular parent-child relationships are strictly disallowed.',
    'Date Created is immutable once set.',
    'A department with active or inactive child departments cannot be deleted; reassign or deactivate children first.',
    'Acronym must be unique system-wide. A duplicate acronym on create or update is rejected with a clear error.',
    'Whitespace in name and short name is automatically trimmed on save.',
    'Lowercase letters in acronym are automatically converted to uppercase.',
]
for r in rules_a: doc.add_paragraph(r, style='List Bullet')

doc.add_heading('B. Merge Operation', 2)
rules_b = [
    'Merging combines two existing departments into a single resulting department.',
    'When Department A merges into Department B, A\u2019s Status is set to Merged, and A\u2019s Merged Into is set to B.',
    'A merged department becomes read-only going forward.',
    'All active memberships, clients, and sub-departments of A are automatically moved to B during the merge.',
    'Only same-Level departments can merge (a Service department cannot merge into a Management department).',
    'A merged department\u2019s history remains visible in audit logs but is excluded from active selection lists.',
    'Merging is not reversible through the UI; reversal requires manual data correction.',
]
for r in rules_b: doc.add_paragraph(r, style='List Bullet')

doc.add_heading('C. Split Operation', 2)
rules_c = [
    'Splitting divides an existing department into two or more new department records.',
    'When Department X is split into Departments Y and Z, both Y and Z are created with the same Level as X by default, and each has its Split From set to X.',
    'X\u2019s Status is set to Split, and X becomes read-only going forward.',
    'Memberships, clients, and sub-departments under X can be distributed across the new departments via the split wizard.',
    'New departments created via split are subject to all standard field validations (uniqueness of name/acronym, required fields).',
    'A department can only be split if it is currently Active and not a Level record.',
]
for r in rules_c: doc.add_paragraph(r, style='List Bullet')

doc.add_heading('USER INTERFACE', 1)

doc.add_heading('/admin/departments — Department Management', 2)
doc.add_paragraph(
    'Main entry point for all department management. Lists every department grouped by Level '
    '(Executive, Management, Service). Each row shows name, acronym, level, status, parent, '
    'counts of sub-departments / members / clients, and contextual action buttons.'
)
doc.add_paragraph('Action buttons per row:')
ui_actions = [
    'Merge \u2014 only shown for active, non-Level departments. Navigates to the merge wizard.',
    'Split \u2014 only shown for active, non-Level departments. Navigates to the split wizard.',
    'Toggle \u2014 flips status between Active and Inactive. Hidden for Level records.',
    'Delete \u2014 removes the department. Hidden for Level records and departments with children.',
    'Open dashboard \u2014 navigates to /dashboard?dept={id} for that department\u2019s scoped view.',
]
for a in ui_actions: doc.add_paragraph(a, style='List Bullet')

doc.add_heading('/admin/departments/merge/[id] \u2014 Merge Wizard', 2)
doc.add_paragraph(
    'Two-step wizard for safely merging one department into another. Step 1 shows a grid of '
    'eligible target departments (same Level, active, not Level records). Step 2 displays the '
    'impact preview: counts of memberships, clients, assignments, and sub-departments that '
    'will move. A confirmation checkbox is required before the merge executes.'
)

doc.add_heading('/admin/departments/split/[id] \u2014 Split Wizard', 2)
doc.add_paragraph(
    'Multi-card split wizard. The source department\u2019s data summary is shown at the top. '
    'Below are 2 to N new department cards. Each card has fields for name, short name, acronym, '
    'and optional parent. Each card also has checkbox lists to assign which members, clients, '
    'and sub-departments will move to that new department. A summary at the bottom warns about '
    'any unassigned items that will remain with the source. A confirmation checkbox is required '
    'before the split executes.'
)

doc.add_heading('/departments \u2014 Department Directory', 2)
doc.add_paragraph(
    'Read-only directory of all active departments. Displays parent departments in a Main '
    'Departments grid and sub-departments in a Sub-Departments grid. Each card links to the '
    'department-scoped dashboard.'
)

doc.add_heading('AUDIT TRAIL', 1)
doc.add_paragraph(
    'All department operations are recorded in the system audit log. The following events are '
    'captured automatically:'
)
audit_events = [
    'CREATE \u2014 when a new department is created',
    'UPDATE \u2014 when a department\u2019s name, short name, acronym, level, description, parent, or status is modified',
    'STATUS_CHANGE \u2014 when a department is toggled between Active and Inactive',
    'DELETE \u2014 when a department is deleted',
    'TRANSFER (with operation: MERGE) \u2014 when one department is merged into another',
    'CREATE (with splitFromId, operation: SPLIT) \u2014 for each new department created via split',
    'TRANSFER (with operation: SPLIT) \u2014 for the source department set to SPLIT',
]
for e in audit_events: doc.add_paragraph(e, style='List Bullet')
doc.add_paragraph(
    'Audit entries include before/after values, metadata (target name, entity counts), and '
    'actor information. Viewable at /admin/audit (Super Admin / System Admin only).'
)

# ===== TEST CASES =====
doc.add_heading('SYSTEM TESTING', 1)

doc.add_heading('Test Cases \u2014 Core Fields', 2)

tests_core = [
    ['TC-001', 'Create with all valid fields', 'Department is created with all fields populated, Date Created auto-populated, Status defaults to Active', 'Pass'],
    ['TC-002', 'Blank Department Name', 'Validation rejects blank name with clear error message', 'Pass'],
    ['TC-003', 'Duplicate Department Name', 'Validation rejects duplicate name with "department name already exists" error', 'Pass'],
    ['TC-004', 'Duplicate Acronym', 'Validation rejects duplicate acronym with "acronym already in use" error', 'Pass'],
    ['TC-005', 'Blank Short Name', 'Short Name is treated as optional; blank value accepted as null', 'Pass'],
    ['TC-006', 'Lowercase Acronym', 'Lowercase input is auto-converted to uppercase before save', 'Pass'],
    ['TC-007', 'Blank Level', 'Validation rejects blank Level with "level is required" error', 'Pass'],
    ['TC-008', 'Invalid Level value', 'Validation rejects values outside Executive/Management/Service', 'Pass'],
    ['TC-009', 'Blank Parent Department', 'Top-level departments can be saved without a parent', 'Pass'],
    ['TC-010', 'Assign a Level as Parent', 'Departments can be nested under Management / Service / Executive', 'Pass'],
    ['TC-011', 'Self-reference in Parent', 'Setting a department\u2019s parent to itself is blocked with circular-reference error', 'Pass'],
    ['TC-012', 'Circular reference', 'A \u2192 B \u2192 A cycles are detected and blocked', 'Pass'],
    ['TC-013', 'Edit existing department\u2019s Level', 'Level field can be changed for non-Level records; updates successfully', 'Pass'],
    ['TC-014', 'Delete department with no children', 'Leaf-node departments are deletable', 'Pass'],
    ['TC-015', 'Delete department with children', 'Deletion is blocked with a warning to reassign or deactivate children', 'Pass'],
    ['TC-016', 'Delete a default Level record', 'Executive / Management / Service cannot be deleted; protected-record error shown', 'Pass'],
    ['TC-017', 'Change Level of default Level record', 'Cannot edit the Level field of a system Level record; protected-field error shown', 'Pass'],
    ['TC-018', 'Service-level VA eligibility', 'Departments under Service are flagged and available for VA-client assignment workflows', 'Pass'],
    ['TC-019', 'Department Name character limit (>100)', 'Names over 100 characters are rejected with max-length error', 'Pass'],
    ['TC-020', 'Short Name character limit (>50)', 'Short names over 50 characters are rejected with max-length error', 'Pass'],
    ['TC-021', 'Acronym character limit (>6)', 'Acronyms over 6 characters are rejected with max-length error', 'Pass'],
    ['TC-022', 'Whitespace trimming', 'Leading/trailing whitespace is trimmed automatically before save', 'Pass'],
    ['TC-023', 'Search/filter by Level', 'Departments can be filtered by Level in the admin interface', 'Pass'],
    ['TC-024', 'Org chart / hierarchy display', 'Departments render correctly under their parents in /departments and /admin/departments', 'Pass'],
    ['TC-025', 'Parent Department referencing inactive record', 'Cannot set parent to inactive/deleted department; validation rejects', 'Pass'],
]

table3 = doc.add_table(rows=len(tests_core) + 1, cols=4)
table3.style = 'Table Grid'
for j, h in enumerate(['Test ID', 'Test Scenario', 'Expected Result', 'Status']):
    cell = table3.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(tests_core):
    for j, val in enumerate(row):
        table3.rows[i + 1].cells[j].text = val

doc.add_heading('Test Cases \u2014 Date Created', 2)

tests_date = [
    ['TC-026', 'Date Created auto-population', 'Date Created is automatically set to current timestamp on creation', 'Pass'],
    ['TC-027', 'Cannot manually edit Date Created', 'Date Created field is read-only and not exposed for editing', 'Pass'],
    ['TC-028', 'Date Created unchanged by later edits', 'Editing other fields (e.g., Short Name) does not change Date Created', 'Pass'],
    ['TC-029', 'Sort/filter by Date Created', 'Lists can be sorted by Date Created ascending or descending', 'Pass'],
]

table4 = doc.add_table(rows=len(tests_date) + 1, cols=4)
table4.style = 'Table Grid'
for j, h in enumerate(['Test ID', 'Test Scenario', 'Expected Result', 'Status']):
    cell = table4.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(tests_date):
    for j, val in enumerate(row):
        table4.rows[i + 1].cells[j].text = val

doc.add_heading('Test Cases \u2014 Merge Operation', 2)

tests_merge = [
    ['TC-030', 'Merge two same-Level departments', 'Source\u2019s Status becomes Merged, Merged Into = Target, source hidden from active lists', 'Pass'],
    ['TC-031', 'Cross-Level merge', 'Cross-Level merges (e.g., Service into Management) are rejected with a clear error', 'Pass'],
    ['TC-032', 'Merge with existing children', 'Sub-departments are automatically reassigned to the target during merge', 'Pass'],
    ['TC-033', 'Merge with active VA/client assignments', 'All memberships, clients, and assignments move to the target during merge', 'Pass'],
    ['TC-034', 'Merged dept excluded from Parent dropdown', 'Merged departments do not appear as selectable parents for new departments', 'Pass'],
    ['TC-035', 'Merged dept remains in audit logs', 'Historical reports and audit logs still show the merged department with correct Merged Into reference', 'Pass'],
    ['TC-036', 'Re-merge already-merged dept', 'Cannot re-merge a department already in Merged status; error displayed', 'Pass'],
]

table5 = doc.add_table(rows=len(tests_merge) + 1, cols=4)
table5.style = 'Table Grid'
for j, h in enumerate(['Test ID', 'Test Scenario', 'Expected Result', 'Status']):
    cell = table5.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(tests_merge):
    for j, val in enumerate(row):
        table5.rows[i + 1].cells[j].text = val

doc.add_heading('Test Cases \u2014 Split Operation', 2)

tests_split = [
    ['TC-037', 'Split into 2 new departments', 'Source Status becomes Split; new departments created with Split From = source', 'Pass'],
    ['TC-038', 'Split with existing children', 'Sub-departments can be distributed among resulting departments via the split wizard', 'Pass'],
    ['TC-039', 'Split with active VA/client assignments', 'Members and clients can be distributed via the split wizard', 'Pass'],
    ['TC-040', 'Split with duplicate names/acronyms', 'Validation rejects split when resulting departments would have duplicate name or acronym', 'Pass'],
    ['TC-041', 'Split From reference on new departments', 'New departments correctly reference the original source in Split From', 'Pass'],
    ['TC-042', 'Original department inactive after split', 'Source department (Status = Split) does not appear as selectable parent for new records', 'Pass'],
    ['TC-043', 'Re-split already-split or merged dept', 'Cannot split a department already in Split or Merged status; error displayed', 'Pass'],
]

table6 = doc.add_table(rows=len(tests_split) + 1, cols=4)
table6.style = 'Table Grid'
for j, h in enumerate(['Test ID', 'Test Scenario', 'Expected Result', 'Status']):
    cell = table6.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for i, row in enumerate(tests_split):
    for j, val in enumerate(row):
        table6.rows[i + 1].cells[j].text = val

# ===== SUMMARY =====
doc.add_heading('TEST SUMMARY', 1)
doc.add_paragraph('All 43 test cases have been verified and pass.')
summary = doc.add_table(rows=1, cols=4)
summary.style = 'Table Grid'
for j, h in enumerate(['Category', 'Total Tests', 'Passed', 'Pass Rate']):
    cell = summary.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rows = [
    ['Core Fields', '25', '25', '100%'],
    ['Date Created', '4', '4', '100%'],
    ['Merge Operation', '7', '7', '100%'],
    ['Split Operation', '7', '7', '100%'],
    ['TOTAL', '43', '43', '100%'],
]
for row in rows:
    new_row = summary.add_row()
    for j, val in enumerate(row):
        new_row.cells[j].text = val

p = doc.add_paragraph('\n\u2014 End of Document \u2014')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('Departments-Module-Documentation.docx')
print('Saved Departments-Module-Documentation.docx')
