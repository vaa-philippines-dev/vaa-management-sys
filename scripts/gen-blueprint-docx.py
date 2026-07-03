"""
VA Management System – Complete Blueprint DOCX Generator
Covers: architecture, database schema, RBAC, all 4 phases,
Google Drive integration, data migration, production deployment,
monitoring, security, roadmap, cost proposal, and risks.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = "F5F5F5"
HEADER_BG = "1F4E79"

# ============================= HELPERS =============================

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = BLUE

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK

def add_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

def add_bold_text(label, value=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK
    if value:
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"\u2022  {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

def add_spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

def shade_cell(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val.get("sz", 4)}" '
            f'w:color="{val.get("color", "CCCCCC")}" w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def make_table(headers, rows, col_widths=None, left_align=False):
    """Create a styled table. col_widths in inches (auto if None)."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        shade_cell(cell, HEADER_BG)
        for edge in ('top', 'bottom', 'left', 'right'):
            set_cell_border(cell, **{edge: {"sz": 4, "color": "CCCCCC"}})

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if left_align else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            if ri % 2 == 1:
                shade_cell(cell, LIGHT_GRAY)
            for edge in ('top', 'bottom', 'left', 'right'):
                set_cell_border(cell, **{edge: {"sz": 4, "color": "CCCCCC"}})

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    add_spacer()
    return table


# ============================= TITLE PAGE =============================

for _ in range(4):
    add_spacer()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VAA Philippines")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Virtual Assistant Management System")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("System Architecture, Database Schema,\nand Production Deployment Blueprint")
r.font.size = Pt(16)
r.font.color.rgb = DARK

add_spacer()
add_spacer()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 2.0 — Full System Design")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_spacer()
add_spacer()
add_text("Prepared for:    VAA Philippines Operations Management")
add_text("Prepared by:     Neil Andre Ibona, Development Team")
add_text("Date:            June 26, 2026")
add_text("Classification:  Internal — Confidential")

doc.add_page_break()

# ============================= 1. EXECUTIVE SUMMARY =============================

add_h1("1. Executive Summary")

add_text(
    "The VAA Philippines Virtual Assistant Management System is a dual-facing "
    "Content Management System (CMS) designed to consolidate all organizational "
    "workflows into a single, scalable platform. It replaces fragmented tools "
    "including Google Sheets with broken import-range formulas, WhatsApp, Google "
    "Chat, and email threads with a unified system supporting manager and "
    "administrator access with granular role-based permissions."
)
add_spacer()
add_text(
    "The system is designed to scale from the current ~3 VA operation to support "
    "700–900 total users with an anticipated daily load of 15–50 active users. "
    "The architecture follows a four-phase development strategy: Department "
    "Hierarchy, VA Workforce, Client Workforce & Service Lines, and Consolidated "
    "Communications (Ticketing)."
)
add_spacer()

add_h3("Key Capabilities")
add_bullet("Dynamic department hierarchy with parent-child relationships")
add_bullet("Distinction between internal staff and Virtual Assistant roles with separate workflows")
add_bullet("Granular RBAC: Super Admin, System Admin, Executive, Department Manager, Contributor, Viewer")
add_bullet("Temporary role assignment for cross-department support without changing primary role")
add_bullet("Tenure tracking with separate records for VA tenure and staff promotion tenure")
add_bullet("Historical audit logging for service reorganizations, position moves, and status changes")
add_bullet("Employment status: Employed, Engaged, Contracted, End of Contract, Transferred, Blacklisted")
add_bullet("Google Drive integration for contracts, clearance records, ID uploads, and portfolios")
add_bullet("Consolidated support ticketing system replacing WhatsApp, Gmail, Google Chat communication")
add_bullet("Comprehensive VA profiles: G-Cash, birthday, emergency contacts, social media, personality traits")
add_bullet("Leave management with distinct approval workflows for staff vs VAs")

# ============================= 2. SYSTEM ARCHITECTURE =============================

add_h1("2. System Architecture Overview")

add_h3("Technology Stack")
make_table(
    ["Layer", "Technology", "Rationale"],
    [
        ["Frontend", "Next.js 16 (App Router) + React 19 + TypeScript", "Server Components for performance, App Router for route groups"],
        ["UI Framework", "Tailwind CSS v4 + ShadCN/ui", "Component library with consistent design tokens"],
        ["Database", "PostgreSQL 16 (via Supabase)", "Relational integrity, JSONB for audit, Row-Level Security, partitioning"],
        ["ORM", "Prisma 7 + pg adapter", "Type-safe queries, migration tooling, schema-as-source-of-truth"],
        ["Auth", "Supabase Auth (Google SSO) + custom RBAC", "Google OAuth with database-level access enforcement"],
        ["File Storage", "Google Drive API (via Service Account)", "Direct Drive integration for contracts, IDs, and documents"],
        ["Caching/Queue", "Redis 7 (Upstash managed)", "Session cache, rate limiting, background job queue"],
        ["Realtime", "Supabase Realtime (WebSockets)", "Live updates for tickets, assignments, and work logs"],
        ["Observability", "OpenTelemetry + Grafana (Loki/Tempo)", "Distributed tracing, structured logging, alerting"],
        ["Hosting", "Vercel Pro (Edge + Serverless)", "Global CDN, zero-config deployment, automatic HTTPS"],
    ],
    col_widths=[1.5, 2.2, 3.8],
    left_align=True,
)

add_h3("High-Level Architecture")
add_text("The system follows a layered architecture with clear separation of concerns:")
add_spacer()
add_bold_text("PRESENTATION LAYER:", " Next.js App Router with route groups for (auth), (dashboard-manager), (dashboard-va), (admin)")
add_bullet("Server Components for data fetching (no client-side API calls for reads)")
add_bullet("Client Components for interactive forms, realtime subscriptions, and file uploads")
add_bullet("Server Actions for all mutations (forms, status changes)")
add_spacer()
add_bold_text("SERVICE LAYER:")
add_bullet("lib/auth.ts — Role-based middleware (requireSuperAdmin, requireManager, requireVA, requireDepartmentAccess)")
add_bullet("lib/google/drive.ts — Google Drive SDK wrapper for folder creation and file uploads")
add_bullet("lib/google/sheets.ts — Google Sheets sync for backward compatibility with legacy data")
add_bullet("lib/audit.ts — Centralized audit trail logger for all state changes")
add_spacer()
add_bold_text("DATA LAYER:")
add_bullet("Prisma ORM with PostgreSQL adapter (lib/prisma.ts singleton)")
add_bullet("PostgreSQL Row-Level Security (RLS) policies enforced at database level")
add_bullet("Supabase Realtime for live data subscriptions on tickets and work logs")
add_spacer()
add_bold_text("EXTERNAL LAYER:")
add_bullet("Google Drive API — Document storage, contract management, file portfolios")
add_bullet("Google Sheets API — Legacy data import/export bridge during migration")
add_bullet("Supabase Auth — Google SSO with session management")

add_h3("Data Flow")
add_text("1. All mutations go through Server Actions → Prisma → PostgreSQL. The UI never calls external APIs directly.")
add_text("2. Google Drive uploads are streamed via API Route Handlers with background processing for large files.")
add_text("3. After successful DB writes, Google Sheets sync is fired as a background promise (fire-and-forget).")
add_text("4. Audit trail entries are created via PostgreSQL triggers for all INSERT/UPDATE/DELETE on critical tables.")
add_text("5. Realtime subscriptions push changes to connected clients for live collaboration features.")

# ============================= 3. DATABASE SCHEMA =============================

add_h1("3. Database Schema — Complete Design")
add_text(
    "The schema is organized into domain groups corresponding to the four implementation "
    "phases. All tables use CUID-based primary keys for distributed unique ID generation "
    "and include created_at/updated_at audit timestamps. Indexes are designed for the "
    "anticipated query patterns at 700+ user scale."
)
add_spacer()

# 3.1 Core Organization
add_h2("3.1 Core Organization (Phase 1 — Foundation)")

add_h3("departments")
add_text("Hierarchical department structure supporting parent-child nesting for organizational chart representation.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", "Unique department identifier"],
        ["name", "String", "UNIQUE, NOT NULL", "Department name (e.g., Amazon Operations, Academy)"],
        ["parent_id", "String?", "FK → departments.id SET NULL", "Parent department for hierarchy (null = top-level)"],
        ["is_parent", "Boolean", "@default(false)", "True if this department acts as a category/group (not assignable directly)"],
        ["description", "String?", "", "Optional description of department function"],
        ["head_id", "String?", "FK → users.id", "Department head / manager responsible"],
        ["sort_order", "Int", "@default(0)", "Display ordering in lists"],
        ["is_active", "Boolean", "@default(true)", "Soft delete / archival flag"],
        ["created_at", "DateTime", "@default(now())", "Record creation timestamp"],
        ["updated_at", "DateTime", "@updatedAt", "Last modification timestamp"],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("positions")
add_text("Job positions/roles within the organization hierarchy. These define the reporting chain and approval workflows.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", "Unique position identifier"],
        ["title", "String", "NOT NULL", "Position title (e.g., Senior Manager, VA, Specialist)"],
        ["reports_to_id", "String?", "FK → positions.id", "Supervisor position in reporting chain"],
        ["department_id", "String?", "FK → departments.id", "Department this position belongs to"],
        ["is_staff_role", "Boolean", "@default(false)", "True if internal staff, false if VA role"],
        ["sort_order", "Int", "@default(0)", "Display ordering"],
        ["is_active", "Boolean", "@default(true)", ""],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("department_memberships")
add_text("Many-to-many join between users and departments, linking to a specific position within each department.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", "Unique membership identifier"],
        ["user_id", "String", "FK → users.id CASCADE", "The user assigned"],
        ["department_id", "String", "FK → departments.id CASCADE", "The department they belong to"],
        ["position_id", "String?", "FK → positions.id", "Their role/position within this department"],
        ["is_primary", "Boolean", "@default(false)", "Primary department for reporting chain"],
        ["started_at", "DateTime", "@default(now())", "When they joined this department"],
        ["ended_at", "DateTime?", "", "When they left (null = current)"],
        ["@@unique([user_id, department_id])", "", "", "One membership per user per department"],
    ],
    col_widths=[1.5, 1.2, 1.5, 3.3],
    left_align=True,
)

# 3.2 Users & Personnel
add_h2("3.2 Users & Personnel (Phase 1–2)")

add_h3("users")
add_text("The central user table redesigned to distinguish staff from VAs and support granular roles.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", "Unique user identifier"],
        ["email", "String", "UNIQUE, NOT NULL", "Login email (Google SSO)"],
        ["first_name", "String", "NOT NULL", "Given name"],
        ["last_name", "String", "NOT NULL", "Family name"],
        ["system_role", "SystemRole enum", "NOT NULL", "SUPER_ADMIN | SYSTEM_ADMIN | EXECUTIVE | DEPT_MANAGER | STAFF | VA"],
        ["user_type", "UserType enum", "NOT NULL", "INTERNAL_STAFF | VIRTUAL_ASSISTANT"],
        ["avatar_url", "String?", "", "Profile photo URL"],
        ["is_active", "Boolean", "@default(true)", "Account status"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("user_profiles")
add_text("Extended profile information for both staff and VAs. One-to-one with users. Includes G-Cash, birthday, emergency contacts, social media, and personality traits for client matching.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["user_id", "String", "UNIQUE, FK → users.id CASCADE", "Link to user record"],
        ["phone", "String?", "", "Primary contact number"],
        ["birth_date", "DateTime?", "", "Date of birth"],
        ["address", "String?", "", "Physical address"],
        ["emergency_contact_name", "String?", "", "Emergency contact person"],
        ["emergency_contact_phone", "String?", "", "Emergency contact number"],
        ["emergency_contact_relation", "String?", "", "Relationship to emergency contact"],
        ["gcash_number", "String?", "", "G-Cash mobile number"],
        ["facebook_url", "String?", "", "Facebook profile link"],
        ["linkedin_url", "String?", "", "LinkedIn profile link"],
        ["personality_traits", "String[]", "", "Traits for client matching (detail-oriented, engaging, etc.)"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.5, 1.2, 1.5, 3.3],
    left_align=True,
)

add_h3("employment_records")
add_text("Tracks employment lifecycle including tenure, contract type, status changes, and department transfers. Separates VA tenure from staff promotion tenure for accurate benefits calculation.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["user_id", "String", "FK → users.id CASCADE", "The employee"],
        ["contract_type", "ContractType enum", "NOT NULL", "REGULAR | PROJECT_BASED | PROBATIONARY"],
        ["employment_status", "EmploymentStatus enum", "NOT NULL", "EMPLOYED | ENGAGED | CONTRACTED | END_OF_CONTRACT | TRANSFERRED | RESIGNED | TERMINATED | BLACKLISTED"],
        ["start_date", "DateTime", "NOT NULL", "Employment start / engagement date"],
        ["end_date", "DateTime?", "", "End date if applicable"],
        ["effective_date", "DateTime", "@default(now())", "When this record became active"],
        ["reason", "String?", "", "Reason for status change (e.g., promotion, resignation)"],
        ["initiated_by", "String?", "FK → users.id", "Who initiated this (HR, service dept)"],
        ["is_current", "Boolean", "@default(true)", "Whether this is the active employment record"],
        ["tenure_days", "Int?", "", "Computed tenure at time of record (for historical accuracy)"],
        ["notes", "String?", "", "Internal notes"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("role_assignments")
add_text("Temporary or additional role assignments for users. Supports Contributor and Viewer roles per module without changing primary system_role.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["user_id", "String", "FK → users.id CASCADE", "The user receiving the role"],
        ["role", "TemporaryRole enum", "NOT NULL", "CONTRIBUTOR | VIEWER | APPROVER"],
        ["module", "String", "NOT NULL", "Target module (e.g., vas, clients, ticketing, reports)"],
        ["department_id", "String?", "FK → departments.id", "Scope limited to specific department"],
        ["granted_by", "String", "FK → users.id", "Admin who assigned this role"],
        ["expires_at", "DateTime?", "", "Auto-expiry for temporary access"],
        ["is_active", "Boolean", "@default(true)", ""],
        ["created_at", "DateTime", "@default(now())", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

# 3.3 VA Workforce
add_h2("3.3 VA Workforce (Phase 2)")

add_h3("va_profiles")
add_text("Enhanced VA profile replacing the simplified version. Adds availability tracking and capacity management.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["user_id", "String", "UNIQUE, FK → users.id CASCADE", "Link to user record"],
        ["hourly_rate", "Decimal?", "", "Billing rate in USD"],
        ["availability_status", "Availability enum", "@default(AVAILABLE)", "AVAILABLE | PARTIALLY_ASSIGNED | FULLY_ASSIGNED | ON_LEAVE | UNAVAILABLE"],
        ["total_capacity_hours", "Decimal?", "@default(40)", "Maximum weekly hours"],
        ["onboarding_folder_url", "String?", "", "Google Drive folder with onboarding docs"],
        ["portfolio_url", "String?", "", "External portfolio or internal Drive link"],
        ["notes", "String?", "", "Manager notes"],
        ["is_active", "Boolean", "@default(true)", "Active in the system"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.3, 1.2, 1.5, 3.5],
    left_align=True,
)

add_h3("va_skills")
add_text("Many-to-many join for VA skill assignments with proficiency levels and years of experience.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["va_profile_id", "String", "FK → va_profiles.id CASCADE", "The VA"],
        ["skill_id", "String", "FK → skills.id CASCADE", "The skill"],
        ["proficiency", "Proficiency enum", "@default(INTERMEDIATE)", "BEGINNER | INTERMEDIATE | ADVANCED | EXPERT"],
        ["years_experience", "Decimal?", "", "Years of experience with this skill"],
        ["is_primary", "Boolean", "@default(false)", "Primary/core skill"],
        ["@@unique([va_profile_id, skill_id])", "", "", "One record per skill per VA"],
    ],
    col_widths=[1.5, 1.2, 1.5, 3.3],
    left_align=True,
)

add_h3("va_documents")
add_text("Document records for contracts, IDs, clearances, and credentials stored in Google Drive. Only metadata in DB.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["va_profile_id", "String", "FK → va_profiles.id CASCADE", ""],
        ["document_type", "DocumentType enum", "NOT NULL", "CONTRACT | GOVERNMENT_ID | NDA | CLEARANCE | CERTIFICATE | ONBOARDING | PERFORMANCE_REVIEW | PORTFOLIO | OTHER"],
        ["file_name", "String", "NOT NULL", "Original filename"],
        ["google_drive_url", "String", "NOT NULL", "Drive file/folder URL"],
        ["mime_type", "String", "", "File MIME type"],
        ["file_size", "Int?", "", "Size in bytes"],
        ["uploaded_by", "String", "FK → users.id", "Who uploaded this document"],
        ["expires_at", "DateTime?", "", "Document expiration (e.g., contract end)"],
        ["notes", "String?", "", ""],
        ["created_at", "DateTime", "@default(now())", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("leave_requests")
add_text("Leave management with approval workflows distinguishing staff vs VA leave processes. Staff follows position reports_to chain; VA routes to department manager.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["user_id", "String", "FK → users.id CASCADE", "Requestor"],
        ["leave_type", "LeaveType enum", "NOT NULL", "VACATION | SICK | EMERGENCY | MATERNITY | PATERNITY | UNPAID | BEREAVEMENT"],
        ["start_date", "DateTime", "NOT NULL", ""],
        ["end_date", "DateTime", "NOT NULL", ""],
        ["total_days", "Decimal", "", "Computed duration"],
        ["reason", "String?", "", ""],
        ["status", "LeaveStatus enum", "@default(PENDING)", "PENDING | APPROVED | REJECTED | CANCELLED"],
        ["approver_id", "String?", "FK → users.id", "Who approved/rejected"],
        ["approved_at", "DateTime?", "", "When decision was made"],
        ["approver_note", "String?", "", ""],
        ["notification_sent", "Boolean", "@default(false)", "Whether VA has been notified"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

# 3.4 Clients & Service Lines
add_h2("3.4 Client Workforce & Service Lines (Phase 3)")

add_h3("clients")
add_text("Client organizations receiving VA services. Enhanced with department assignment and onboarding folders.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["name", "String", "NOT NULL", "Client company/brand name"],
        ["contact_name", "String?", "", "Primary contact person"],
        ["contact_email", "String?", "", "Primary contact email"],
        ["contact_phone", "String?", "", "Primary contact phone"],
        ["platform", "ClientPlatform enum", "@default(MULTI)", "AMAZON | WALMART | TIKTOK_SHOP | SHOPIFY | MULTI"],
        ["industry", "String?", "", "Industry/niche"],
        ["timezone", "String?", "@default('America/New_York')", "Client timezone for scheduling"],
        ["notes", "String?", "", "Internal notes"],
        ["is_active", "Boolean", "@default(true)", ""],
        ["manager_id", "String?", "FK → users.id", "Account manager"],
        ["department_id", "String?", "FK → departments.id", "Servicing department"],
        ["onboarding_folder_url", "String?", "", "Google Drive client onboarding folder"],
        ["required_skills", "String[]", "", "Required services for VA matching"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("assignments")
add_text("Matching VAs to clients with agreed scope and tracking. Skill requirements stored for the matching algorithm.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["type", "AssignmentType enum", "NOT NULL", "REGULAR | PROJECT"],
        ["status", "AssignmentStatus enum", "@default(ACTIVE)", "ACTIVE | PAUSED | COMPLETED | CANCELLED | ON_HOLD"],
        ["agreed_hours", "Decimal", "NOT NULL", "Total agreed hours for assignment"],
        ["monthly_hours", "Decimal?", "", "Monthly hour target for retainer-type"],
        ["start_date", "DateTime", "NOT NULL", ""],
        ["end_date", "DateTime?", "", "Project end date or contract end"],
        ["notes", "String?", "", ""],
        ["va_profile_id", "String", "FK → va_profiles.id CASCADE", "Assigned VA"],
        ["client_id", "String", "FK → clients.id CASCADE", "Receiving client"],
        ["skill_requirements", "String[]", "", "Skills required for this assignment"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("work_logs")
add_text("Daily hour entries recorded by VAs per assignment.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["work_date", "DateTime", "NOT NULL", "Date work was performed"],
        ["hours", "Decimal", "NOT NULL", "Hours logged"],
        ["description", "String?", "", "Description of work done"],
        ["va_profile_id", "String", "FK → va_profiles.id CASCADE", ""],
        ["assignment_id", "String", "FK → assignments.id CASCADE", ""],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

# 3.5 Ticketing
add_h2("3.5 Consolidated Communications & Ticketing (Phase 4)")

add_h3("tickets")
add_text("Centralized support ticket system replacing WhatsApp, email, and chat threads. Auto-assigns to department based on category and client.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["ticket_number", "String", "UNIQUE", "Auto-generated (e.g., TKT-2026-0001)"],
        ["title", "String", "NOT NULL", "Brief summary"],
        ["description", "String?", "", "Detailed description"],
        ["category", "TicketCategory enum", "NOT NULL", "TECHNICAL | HR | CLIENT | VA_SUPPORT | GENERAL"],
        ["priority", "Priority enum", "@default(MEDIUM)", "LOW | MEDIUM | HIGH | URGENT"],
        ["status", "TicketStatus enum", "@default(OPEN)", "OPEN | IN_PROGRESS | WAITING_ON_CLIENT | RESOLVED | CLOSED"],
        ["source", "TicketSource enum", "@default(INTERNAL)", "INTERNAL | EMAIL | WHATSAPP | CLIENT_PORTAL"],
        ["created_by", "String", "FK → users.id", "Ticket creator"],
        ["assigned_to", "String?", "FK → users.id", "Assignee"],
        ["department_id", "String?", "FK → departments.id", "Responsible department"],
        ["client_id", "String?", "FK → clients.id", "Related client if applicable"],
        ["resolved_at", "DateTime?", "", "Resolution timestamp"],
        ["created_at", "DateTime", "@default(now())", ""],
        ["updated_at", "DateTime", "@updatedAt", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

add_h3("ticket_conversations")
add_text("Threaded conversation history within tickets. Supports internal notes not visible to clients.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["ticket_id", "String", "FK → tickets.id CASCADE", ""],
        ["user_id", "String", "FK → users.id", "Message author"],
        ["message", "String", "NOT NULL", "Message content"],
        ["is_internal_note", "Boolean", "@default(false)", "Internal notes not visible to clients"],
        ["attachments", "String[]", "", "Google Drive URLs of attachments"],
        ["created_at", "DateTime", "@default(now())", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

# 3.6 Audit Trail
add_h2("3.6 Audit Trail & Historical Logging")

add_h3("audit_logs")
add_text("Immutable audit trail for all system changes. Supports historical tracking for service reorganizations, position moves, and status transitions. Uses JSONB for old/new value snapshots.")
make_table(
    ["Column", "Type", "Constraints", "Description"],
    [
        ["id", "String (cuid)", "PK", ""],
        ["actor_id", "String", "FK → users.id", "Who performed the action"],
        ["action", "AuditAction enum", "NOT NULL", "CREATE | UPDATE | DELETE | STATUS_CHANGE | TRANSFER | APPROVE | REJECT"],
        ["entity_type", "String", "NOT NULL", "Affected model (e.g., User, Department, Assignment)"],
        ["entity_id", "String", "NOT NULL", "Affected record ID"],
        ["old_values", "Json?", "", "Previous state snapshot"],
        ["new_values", "Json?", "", "New state snapshot"],
        ["metadata", "Json?", "", "Additional context (IP, user agent, etc.)"],
        ["department_id", "String?", "FK → departments.id", "Scope of the change"],
        ["created_at", "DateTime", "@default(now())", ""],
    ],
    col_widths=[1.2, 1.2, 1.5, 3.6],
    left_align=True,
)

# ============================= 4. ACCESS CONTROL & RBAC =============================

add_h1("4. Access Control & Role-Based Permissions")

add_h3("System Roles (system_role enum)")
make_table(
    ["Role", "Permissions", "Who Gets This"],
    [
        ["SUPER_ADMIN", "Full system access: all modules, configuration, user management, department creation, audit logs", "VAA Philippines IT/Operations Head"],
        ["SYSTEM_ADMIN", "User management, department configuration, role assignment, but no data deletion or config changes", "Senior Admin Staff"],
        ["EXECUTIVE", "Read-only access to all modules, reports, and dashboards. Cannot modify data.", "CEO, CMO, COO, Senior Managers"],
        ["DEPT_MANAGER", "Full access within their department: VAs, clients, assignments, tickets. Cross-dept only if granted.", "Department Heads, Service Leads"],
        ["STAFF", "Read/write within their department scope. Can create tickets, log work, view own data.", "Internal Staff Employees"],
        ["VA", "Read/write own profile, assignments, work logs, tickets. Cannot view other VAs or financial data.", "Virtual Assistants"],
    ],
    col_widths=[1.5, 3.5, 2.5],
    left_align=True,
)

add_h3("Temporary Roles (role_assignments)")
add_text("Temporary role elevations allow staffing flexibility without changing primary permissions:")
add_bullet("CONTRIBUTOR — Can create and edit within a specific module and department scope")
add_bullet("VIEWER — Read-only access to a specific module and department scope")
add_bullet("APPROVER — Can approve leave requests, status changes within scope")
add_text("These roles automatically expire based on expires_at or are manually revoked. The audit log records all temporary role grants and revocations.")
add_spacer()

add_h3("View-As Feature")
add_text("Super Admin and System Admin users can preview the system from another role's perspective without logging out. This is a UI-layer toggle that applies the target role's RLS policies for the session.")

# ============================= 5. PHASE 1: DEPARTMENT HIERARCHY =============================

add_h1("5. Phase 1: Department Hierarchy")
add_text("The foundation of the entire system. Establishes the organizational structure that all other modules depend on.")
add_spacer()
add_h3("Implementation Scope")
add_bullet("Create departments table with self-referencing parent_id for hierarchy")
add_bullet("Create positions table with reports_to chain")
add_bullet("Create department_memberships for many-to-many user-department relationships")
add_bullet("Migrate existing users to the new schema with proper department assignments")
add_bullet("Build department management UI (admin-only CRUD)")
add_bullet("Organizational chart visualization component")
add_spacer()
add_bold_text("PostgreSQL RLS Policy: ", "Managers can read/write within their department scope. Super/System Admins have full access.")
add_spacer()

# ============================= 6. PHASE 2: VA WORKFORCE =============================

add_h1("6. Phase 2: VA Workforce Management")
add_text("Comprehensive VA management with enhanced profiles, skill matching, document storage, and leave management.")
add_spacer()
add_h3("Implementation Scope")
add_bullet("Enhanced VA profiles with full personal data (G-Cash, birthday, emergency contacts, social media)")
add_bullet("Skills management with proficiency levels and years of experience")
add_bullet("Employment records with tenure tracking (separate VA tenure vs staff promotion tenure)")
add_bullet("Document management integrated with Google Drive (contracts, IDs, clearances)")
add_bullet("Leave request workflow with department-level approval chain")
add_bullet("Availability status tracking and capacity management")
add_bullet("Blacklist management for restricted rehire tracking")
add_bullet("Personality trait capture for future client matching")
add_spacer()
add_bold_text("Key Workflow: ", "Leave requests for VAs route to their department manager. Staff leave follows the position reports_to chain. The approver receives an email/push notification.")

doc.add_page_break()

# ============================= 7. PHASE 3: CLIENTS =============================

add_h1("7. Phase 3: Client Workforce & Service Lines")
add_text("Client management, service line configuration, and the VA-client assignment matching engine.")
add_spacer()
add_h3("Implementation Scope")
add_bullet("Enhanced client profiles with onboarding folders and department assignment")
add_bullet("Service line modules mapped to departments (e.g., Amazon Services, Shopify Support)")
add_bullet("Skill-based VA-to-client matching algorithm (suggest top 3 VAs per client skill requirements)")
add_bullet("Assignment lifecycle management with status workflow (Active → Paused → Completed/Cancelled)")
add_bullet("Google Drive integration for client onboarding document storage")
add_bullet("Historical logging for service/position moves between departments")
add_bullet("Reporting and dashboards per client (hours, utilization, assignment status)")
add_spacer()

# ============================= 8. PHASE 4: TICKETING =============================

add_h1("8. Phase 4: Consolidated Communications & Ticketing")
add_text("Replaces fragmented communication channels (WhatsApp, Google Chat, email) with a centralized support ticket system.")
add_spacer()
add_h3("Implementation Scope")
add_bullet("Ticket creation with categories (Technical, HR, Client, VA Support, General)")
add_bullet("Threaded conversation history with file attachments")
add_bullet("Ticket routing: auto-assign to department based on category and client")
add_bullet("Internal notes vs client-facing messages")
add_bullet("WhatsApp/email integration hooks for inbound ticket creation (future)")
add_bullet("SLA tracking and escalation (auto-escalate tickets exceeding response time)")
add_bullet("Ticket-to-assignment linking (connect support tickets to VA assignments)")
add_bullet("Real-time notifications via Supabase Realtime when tickets are updated")
add_spacer()

# ============================= 9. GOOGLE DRIVE =============================

add_h1("9. Google Drive & Document Management Integration")
add_text("The system uses Google Drive API via a Service Account for all document storage. Files are never stored on the application server.")
add_spacer()
add_h3("Integration Architecture")
add_bullet("Service Account with Drive scope: https://www.googleapis.com/auth/drive.file")
add_bullet("Parent folder structure: /VAA Philippines/{Department}/{VA Name}/{Document Type}/")
add_bullet("Automatic folder creation on VA profile creation")
add_bullet("Direct file upload via API Route Handler → Drive SDK (streamed, not buffered)")
add_bullet("2-second average upload delay (acceptable per manager discussion)")
add_bullet("Only Google Drive URLs stored in the database, never file binaries")
add_bullet("Mock implementation available for local development (lib/google/drive.ts)")
add_spacer()
add_h3("Document Types & Categories")
add_text("CONTRACT | GOVERNMENT_ID | NDA | CLEARANCE | CERTIFICATE | ONBOARDING | PERFORMANCE_REVIEW | PORTFOLIO | OTHER")

# ============================= 10. DATA MIGRATION =============================

add_h1("10. Data Migration Strategy")
add_text("The master list from the legacy Google Sheet will be imported in two phases to maintain data integrity.")
add_spacer()
add_h3("Phase A: Department Structure")
add_bullet("Extract unique department names from the master list")
add_bullet("Map parent-child relationships based on organizational chart")
add_bullet("Create departments with proper hierarchy in the new schema")
add_bullet("Validate against Ian's department data shared via WhatsApp")
add_spacer()
add_h3("Phase B: Personnel & VA Data")
add_bullet("Import user records with first_name, last_name, email from legacy sheet")
add_bullet("Create employment records matching existing tenure data")
add_bullet("Create VA profiles with hourly rates, skills mapping")
add_bullet("Link users to departments via department_memberships")
add_bullet("Generate audit trail entries noting data migration source")
add_spacer()
add_h3("Data Validation")
add_bullet("Cross-reference row counts between legacy sheet and database")
add_bullet("Spot-check 10% of records for accuracy")
add_bullet("Keep legacy sheet as read-only reference during transition period (2 weeks)")

# ============================= 11. PRODUCTION DEPLOYMENT =============================

add_h1("11. Production Deployment Architecture")

add_h3("Infrastructure")
make_table(
    ["Component", "Production", "Development"],
    [
        ["Web Hosting", "Vercel Pro (Edge + Serverless)", "Vercel Hobby / localhost:3000"],
        ["Database", "Supabase PostgreSQL (managed)", "Supabase free tier / local PostgreSQL"],
        ["Auth", "Supabase Auth (Google SSO)", "Supabase local / mock"],
        ["File Storage", "Google Drive API (Service Account)", "Local mock (.google-mock/)"],
        ["Cache/Queue", "Upstash Redis", "Local Redis / ioredis-mock"],
        ["Monitoring", "Grafana Cloud (Loki + Tempo)", "Console logs"],
        ["CI/CD", "GitHub Actions → Vercel deploy", "Local dev server"],
    ],
    col_widths=[1.8, 3.0, 2.7],
    left_align=True,
)

add_h3("Deployment Workflow")
add_text("1. Push to main branch triggers GitHub Actions workflow")
add_text("2. Prisma migration runs against production database")
add_text("3. Build: next build with production optimizations")
add_text("4. Deploy to Vercel with automatic preview deployments on PR")
add_text("5. Health check: GET /api/health verifies DB connectivity and Drive API access")
add_spacer()

# ============================= 12. MONITORING =============================

add_h1("12. Monitoring, Logging & Observability")
add_bullet("Structured logging via console (JSON format) forwarded to Grafana Loki")
add_bullet("OpenTelemetry traces for all Server Actions (createDepartment, processTicket, etc.)")
add_bullet("Supabase health metrics: connection pool, query performance, RLS policy hits")
add_bullet("Error tracking: Unhandled errors caught by Next.js error boundary and logged with stack traces")
add_bullet("Performance monitoring: Vercel Analytics for Web Vitals, custom instrumentation for DB query times")
add_bullet("Alerting: Grafana alerts on error rate > 1%, p95 latency > 2s, DB connection pool exhaustion")
add_spacer()

# ============================= 13. SECURITY =============================

add_h1("13. Security & Compliance")
add_bullet("Google SSO with domain-restricted access (only @vaa.ph emails)")
add_bullet("Row-Level Security on all tables enforced at PostgreSQL level")
add_bullet("Column-level masking for financial data (salary, rates) not visible to VA role")
add_bullet("All Google Drive files stored under Service Account ownership, not user accounts")
add_bullet("Audit trail for every CREATE/UPDATE/DELETE on critical tables (immutable records)")
add_bullet("Rate limiting on login attempts and API routes via Upstash Redis")
add_bullet("HTTPS enforced via Vercel automatic SSL")
add_bullet("Environment secrets stored in Vercel Environment Variables (never in source)")
add_bullet("CORS restricted to the application origin only")
add_spacer()

# ============================= 14. COST PROPOSAL =============================

add_h1("14. Infrastructure Cost Proposal")
add_text(
    "The VAA Philippines VA Management System is being upgraded from a basic "
    "prototype to a production-grade platform capable of scaling from 3 VAs to "
    "700+ VAs. This section outlines the infrastructure cost at each growth phase "
    "and compares fully-managed vs. self-hosted strategies."
)
add_spacer()
add_bold_text("Key takeaway: ", "Starting costs are minimal ($57/mo at 50 VAs). The architecture is designed so you can start fully managed and seamlessly transition to self-hosted as you exceed 200 VAs, without rewriting any code.")
add_spacer()

add_h3("14.1 Monthly Cost Comparison by VA Tier")
make_table(
    ["Category", "50 VAs", "200 VAs", "700 VAs"],
    [
        ["PostgreSQL + Auth", "$25", "$150", "$600"],
        ["Web Hosting (Vercel)", "$20", "$20", "$20"],
        ["Redis Stream (queue)", "$5", "$12", "$25"],
        ["Go Workers (ingestion)", "$5", "$15", "$25"],
        ["Object Storage (screenshots)", "$1.50", "$5", "$15"],
        ["Observability", "$0", "$50", "$150"],
        ["Fully Managed Total", "$56.50", "$252", "$835"],
        ["Self-Hosted Total", "$56.50", "$155", "$155"],
    ],
    col_widths=[3.0, 1.5, 1.5, 1.5],
)
add_text("New infrastructure cost to serve 50 VAs is only $11.50/mo on top of current spend.")
add_spacer()

add_h3("14.2 At 700 VAs: Fully Managed vs. Self-Hosted")
add_bold_text("700 VA monthly data volume:")
add_bullet("Time ticks: 6 million per day (every 10 seconds per VA)")
add_bullet("Database rows: 132 million per month")
add_bullet("Screenshots: 201,600 per day = ~880 GB per month")
add_bullet("PostgreSQL storage: ~500 GB after 3 months")
add_spacer()
add_h3("Side-by-Side Comparison")
make_table(
    ["Resource", "Fully Managed", "Self-Hosted Hybrid", "Monthly Savings"],
    [
        ["PostgreSQL", "Supabase Team — $599/mo", "Hetzner CX62 (8vCPU, 32GB) — $80/mo", "$519"],
        ["Web Hosting", "Vercel Pro — $20/mo", "Vercel Pro — $20/mo", "$0"],
        ["Redis", "Upstash 4GB — $25/mo", "Self-hosted on DB VPS — $0", "$25"],
        ["Workers", "Railway 4 instances — $25/mo", "Hetzner CX42 (4vCPU, 8GB) — $40/mo", "-$15"],
        ["Screenshots", "Cloudflare R2 — $15/mo", "Cloudflare R2 — $15/mo", "$0"],
        ["Logs + Traces", "Grafana Cloud Pro — $150/mo", "Self-hosted on VPS — $0", "$150"],
        ["Total Monthly", "$834/mo", "$155/mo", "-$679/mo"],
        ["Total Yearly", "$10,008/yr", "$1,860/yr", "-$8,148/yr"],
    ],
    col_widths=[1.8, 2.5, 2.4, 0.8],
)

add_h3("Self-Hosted Hardware Specs")
add_bold_text("VPS 1 — Database Server (Hetzner CX62) — $80/mo")
add_bullet("8 vCPU (AMD EPYC), 32 GB RAM, 500 GB NVMe SSD, 20 TB traffic")
add_bullet("Runs: PostgreSQL 16 + Redis + Grafana stack")
add_spacer()
add_bold_text("VPS 2 — Worker Server (Hetzner CX42) — $40/mo")
add_bullet("4 vCPU (AMD EPYC), 8 GB RAM, 80 GB NVMe SSD, 20 TB traffic")
add_bullet("Runs: Go ingestion workers + OpenTelemetry collector")
add_spacer()
add_text("Total infrastructure: $120/mo + $35/mo (R2 + Vercel) = $155/mo")
add_spacer()

add_h3("14.3 Yearly Cost Projection")
make_table(
    ["VA Tier", "Fully Managed (Annual)", "Self-Hosted (Annual)", "Savings"],
    [
        ["50", "$678", "$678", "$0"],
        ["200", "$3,024", "$1,860", "$1,164"],
        ["700", "$10,008", "$1,860", "$8,148"],
    ],
    col_widths=[2.0, 2.0, 2.0, 1.5],
)
add_text("At 700 VAs, self-hosting saves $8,148 per year — enough to cover a part-time DevOps contractor for monitoring.")
add_spacer()

add_h3("14.4 Technology Stack")
make_table(
    ["Layer", "Technology", "Why"],
    [
        ["Workers", "Go", "Sub-ms latency, native concurrency for high-throughput ingestion"],
        ["Web App", "Next.js 16 + TypeScript", "Already built on this stack"],
        ["Database", "PostgreSQL 16 + partitioning", "Relational integrity, JSONB for audit, Row-Level Security"],
        ["Queue", "Redis 7 Streams", "Single tool for queue + cache. No Kafka needed at 6M msg/day"],
        ["Object Storage", "Cloudflare R2", "Zero egress fees vs AWS S3 ($90/TB). Saves $79/mo at 700 VAs"],
        ["Auth", "Supabase Auth + custom RBAC", "Google SSO already integrated"],
        ["Observability", "OpenTelemetry → Grafana", "Correlation-ID tracing across all services"],
    ],
    col_widths=[1.5, 2.2, 3.8],
    left_align=True,
)

add_h3("Why NOT These Alternatives")
make_table(
    ["Rejected", "Reason"],
    [
        ["AWS S3 for screenshots", "$90/TB egress at 700 VAs — R2 is $0/TB"],
        ["RabbitMQ / Kafka", "Overkill for 6M msg/day. Redis Streams handles this"],
        ["Celery (Python workers)", "Python GIL limits throughput vs Go"],
        ["AWS RDS Aurora", "3-4x more expensive than Hetzner VPS for same performance"],
    ],
    col_widths=[2.5, 5.0],
    left_align=True,
)

add_h3("14.5 Database Growth Projection")
make_table(
    ["VA Count", "Ticks/Day", "Ticks/Month", "DB Storage (3mo)", "Screenshots/Month", "R2 Storage (3mo)"],
    [
        ["50", "432,000", "9.5M", "~35 GB", "14,400", "~25 GB"],
        ["200", "1.7M", "38M", "~140 GB", "57,600", "~100 GB"],
        ["700", "6M", "132M", "~500 GB", "201,600", "~880 GB"],
    ],
    col_widths=[1.0, 1.2, 1.2, 1.3, 1.4, 1.3],
)

add_h3("14.6 Phased Rollout Plan")
make_table(
    ["Phase", "VAs", "Stack", "Monthly Cost", "Timeline"],
    [
        ["Now", "3", "Current (Supabase + Vercel)", "$45", "Already running"],
        ["1", "≤50", "Add Redis + Workers + R2", "~$57", "Week 1-2"],
        ["2", "50-200", "Upgrade Supabase to Scale", "~$252", "Month 1-2"],
        ["3", "200+", "Migrate to self-hosted PG", "~$155", "Month 3+"],
    ],
    col_widths=[1.0, 0.8, 2.5, 1.3, 1.7],
)
add_text("The migration from Supabase to self-hosted requires zero code changes. Prisma abstracts the database connection. It is a one-evening operation: pg_dump from Supabase → psql restore to Hetzner VPS → update DATABASE_URL env var → deploy.")
add_spacer()

add_h3("14.7 One-Time Development Costs")
make_table(
    ["Component", "Effort", "Description"],
    [
        ["Go ingestion worker", "16 hrs", "Redis Streams to PostgreSQL batch insert"],
        ["Tracker client (IndexedDB + presigned URLs)", "8 hrs", "Browser extension or Electron app"],
        ["Audit triggers + RLS policies", "4 hrs", "SQL triggers and row-level security"],
        ["RBAC middleware extension", "4 hrs", "Route guards + financial column masking"],
        ["Total", "32 hrs (one sprint)", ""],
    ],
    col_widths=[3.0, 1.0, 3.5],
    left_align=True,
)

add_h3("14.8 Key Risks & Mitigations")
make_table(
    ["Risk", "Impact", "Mitigation"],
    [
        ["PostgreSQL slow at 100M+ rows/mo", "Queries take minutes", "Monthly partitioning + composite indexes on (va_profile_id, recorded_at)"],
        ["Redis data loss on crash", "Lose up to 5 seconds of ticks", "AOF persistence + client-side IndexedDB replay"],
        ["Worker crash mid-batch", "Duplicate rows", "Dedup via client_batch_id in Redis TTL set"],
        ["Philippines ISP outages", "Lost time tracking", "IndexedDB local queue with sync protocol on reconnect"],
        ["Vercel cold starts on API", "500ms latency on tick ingestion", "Edge functions for POST /api/ticks (<50ms)"],
        ["Manager sees billing data", "Policy violation", "PostgreSQL column-level VIEW masking at DB level"],
    ],
    col_widths=[2.2, 2.0, 3.3],
    left_align=True,
)

doc.add_page_break()

# ============================= 15. DEVELOPMENT ROADMAP =============================

add_h1("15. Development Roadmap")
make_table(
    ["Phase", "Timeline", "Deliverables", "Dependencies"],
    [
        ["0 (Current)", "Complete", "Auth, Basic VA, Client, Assignment, Work Log CRUD", "None"],
        ["0.5 (Schema v2)", "Week 1", "Full schema migration plan, Prisma schema v2.0", "Phase 0 complete"],
        ["1 (Departments)", "Week 1-2", "Departments, Positions, Memberships, Org Chart UI", "Schema v2"],
        ["2 (VA Workforce)", "Week 2-4", "Enhanced VA profiles, Skills v2, Employment records, Leave, Documents", "Phase 1"],
        ["3 (Clients & Services)", "Week 4-6", "Enhanced Clients, Service Lines, Skill Matching, Assignment v2", "Phase 2"],
        ["4 (Ticketing)", "Week 6-8", "Ticket system, Conversations, SLA tracking, Notifications", "Phase 3"],
        ["5 (Polish)", "Week 8-9", "Reports, Dashboards, Loading indicators, Performance optimization", "Phase 4"],
        ["6 (Deploy)", "Week 9-10", "Production deployment, Data migration, UAT, Go-live", "Phase 5"],
    ],
    col_widths=[1.3, 1.3, 3.0, 1.7],
    left_align=True,
)

doc.add_page_break()

# ============================= 16. APPENDIX: ENUMS =============================

add_h1("16. Appendix: Complete Enumerations")

add_h3("New Enums (Schema v2)")
make_table(
    ["Enum Name", "Values"],
    [
        ["SystemRole", "SUPER_ADMIN | SYSTEM_ADMIN | EXECUTIVE | DEPT_MANAGER | STAFF | VA"],
        ["UserType", "INTERNAL_STAFF | VIRTUAL_ASSISTANT"],
        ["TemporaryRole", "CONTRIBUTOR | VIEWER | APPROVER"],
        ["ContractType", "REGULAR | PROJECT_BASED | PROBATIONARY"],
        ["EmploymentStatus", "EMPLOYED | ENGAGED | CONTRACTED | END_OF_CONTRACT | TRANSFERRED | RESIGNED | TERMINATED | BLACKLISTED"],
        ["Availability", "AVAILABLE | PARTIALLY_ASSIGNED | FULLY_ASSIGNED | ON_LEAVE | UNAVAILABLE"],
        ["Proficiency", "BEGINNER | INTERMEDIATE | ADVANCED | EXPERT"],
        ["LeaveType", "VACATION | SICK | EMERGENCY | MATERNITY | PATERNITY | UNPAID | BEREAVEMENT"],
        ["LeaveStatus", "PENDING | APPROVED | REJECTED | CANCELLED"],
        ["TicketCategory", "TECHNICAL | HR | CLIENT | VA_SUPPORT | GENERAL"],
        ["TicketStatus", "OPEN | IN_PROGRESS | WAITING_ON_CLIENT | RESOLVED | CLOSED"],
        ["TicketSource", "INTERNAL | EMAIL | WHATSAPP | CLIENT_PORTAL"],
        ["DocumentType", "CONTRACT | GOVERNMENT_ID | NDA | CLEARANCE | CERTIFICATE | ONBOARDING | PERFORMANCE_REVIEW | PORTFOLIO | OTHER"],
        ["AuditAction", "CREATE | UPDATE | DELETE | STATUS_CHANGE | TRANSFER | APPROVE | REJECT"],
    ],
    col_widths=[2.2, 5.3],
    left_align=True,
)

add_h3("Existing Enums (Retained from v1)")
make_table(
    ["Enum Name", "Values"],
    [
        ["AssignmentType", "REGULAR | PROJECT"],
        ["AssignmentStatus", "ACTIVE | PAUSED | COMPLETED | CANCELLED | ON_HOLD"],
        ["SkillCategory", "STANDARD | UPSKILL | SPECIAL"],
        ["ClientPlatform", "AMAZON | WALMART | TIKTOK_SHOP | SHOPIFY | MULTI"],
        ["Priority", "LOW | MEDIUM | HIGH | URGENT"],
    ],
    col_widths=[2.2, 5.3],
    left_align=True,
)

add_spacer()
add_spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— End of Document —")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================= SAVE =============================

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "VA-Management-System-Blueprint.docx")
doc.save(out_path)
print(f"Created: {os.path.abspath(out_path)}")
